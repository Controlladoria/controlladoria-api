"""
Week 2: Structured document processor
Extracts structured JSON data from financial documents
"""

import base64
import hashlib
import io
import json
import logging
import os
import re
import threading
import time
from collections import defaultdict
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
import google.generativeai as genai
from pdf2image import convert_from_path
from PIL import Image

from ai_key_pool import KeyPool
from models import (
    CategorySummary,
    DateRangeSummary,
    FinancialDocument,
    Transaction,
    TransactionLedger,
)
from i18n_errors import translate_error, translate_ai_error, translate_validation_error

# Redis for caching (optional)
try:
    import redis

    REDIS_AVAILABLE = True
except ImportError:
    REDIS_AVAILABLE = False
    logging.warning("Redis not available - AI caching disabled")

# XML processing
try:
    import lxml.etree as ET
    import xmltodict

    XML_AVAILABLE = True
except ImportError:
    XML_AVAILABLE = False
    logging.warning("XML processing libraries (xmltodict, lxml) not available")

load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

AI_PROVIDER = os.getenv("AI_PROVIDER", "openai")


class StructuredDocumentProcessor:
    """Processes documents and extracts structured financial data"""

    def __init__(self):
        # AI_PROVIDER supports comma-separated values: "gemini,nova,openai"
        # All listed providers are used as co-primaries (round-robin across all).
        # Unlisted providers with registered keys are still used as failover.
        raw_provider = AI_PROVIDER or "gemini"
        self.ai_providers = [p.strip() for p in raw_provider.split(",") if p.strip()]
        self.ai_provider = self.ai_providers[0]  # backward compat: first is "primary"
        logger.info(
            f"Initializing StructuredDocumentProcessor with AI provider(s): {', '.join(self.ai_providers)}"
        )

        # Load model configuration (2026 latest models)
        self.openai_model = os.getenv("OPENAI_MODEL", "gpt-5.4-nano")
        self.gemini_model = os.getenv("GEMINI_MODEL", "gemini-flash-lite-latest")
        self.nova_model = os.getenv("NOVA_MODEL", "us.amazon.nova-lite-v2:0")
        self.nova_region = os.getenv("NOVA_REGION", os.getenv("AWS_REGION", "us-east-2"))

        # Caching configuration (disabled by default - requires Redis)
        self.enable_cache = os.getenv("ENABLE_AI_CACHE", "False").lower() == "true"
        self.cache_ttl = int(os.getenv("AI_CACHE_TTL", "86400"))  # 24 hours

        # Retry configuration
        self.max_retries = int(os.getenv("AI_MAX_RETRIES", "3"))
        self.retry_delay = int(os.getenv("AI_RETRY_DELAY", "1"))
        self.timeout = int(os.getenv("AI_TIMEOUT", "60"))

        # Failover configuration
        self.ai_failover_enabled = os.getenv("AI_FAILOVER_ENABLED", "true").lower() == "true"

        # Initialize Redis cache if available and enabled
        self.cache = None
        if self.enable_cache and REDIS_AVAILABLE:
            try:
                redis_url = os.getenv("REDIS_URL", "redis://localhost:6379/0")
                self.cache = redis.from_url(redis_url, decode_responses=True)
                self.cache.ping()  # Test connection
                logger.info(f"✓ Redis cache enabled (TTL: {self.cache_ttl}s)")
            except Exception as e:
                logger.warning(f"⚠️  Redis connection failed, caching disabled: {e}")
                self.cache = None

        # --- Key Pool Setup (round-robin multi-key) ---
        unhealthy_threshold = int(os.getenv("AI_KEY_UNHEALTHY_THRESHOLD", "3"))
        recovery_seconds = int(os.getenv("AI_KEY_RECOVERY_SECONDS", "300"))
        self.key_pool = KeyPool(
            unhealthy_threshold=unhealthy_threshold,
            recovery_seconds=recovery_seconds,
        )

        # Collect OpenAI keys (multi-key comma-separated, fallback to single key)
        openai_keys_str = os.getenv("OPENAI_API_KEYS", "")
        openai_keys = [k.strip() for k in openai_keys_str.split(",") if k.strip()] if openai_keys_str else []
        if not openai_keys:
            single_key = os.getenv("OPENAI_API_KEY", "")
            if single_key:
                openai_keys = [single_key]
        if openai_keys:
            self.key_pool.register_keys("openai", openai_keys, self.openai_model)

        # Collect Gemini keys (multi-key comma-separated, fallback to single key)
        gemini_keys_str = os.getenv("GEMINI_API_KEYS", "")
        gemini_keys = [k.strip() for k in gemini_keys_str.split(",") if k.strip()] if gemini_keys_str else []
        if not gemini_keys:
            single_key = os.getenv("GEMINI_API_KEY", "")
            if single_key:
                gemini_keys = [single_key]
        if gemini_keys:
            self.key_pool.register_keys("gemini", gemini_keys, self.gemini_model)

        # Register Nova (Amazon Bedrock) if AWS credentials exist
        if os.getenv("AWS_ACCESS_KEY_ID") or os.getenv("AWS_PROFILE"):
            self.key_pool.register_keys("nova", ["iam-credentials"], self.nova_model)

        # Validate: at least one of the configured providers must have keys
        configured_with_keys = [p for p in self.ai_providers if self.key_pool.has_provider(p)]
        if not configured_with_keys:
            raise ValueError(
                f"No API keys configured for any of the AI providers: {self.ai_providers}. "
                f"Set GEMINI_API_KEY(S), AWS credentials, or OPENAI_API_KEY(S)."
            )

        # Build provider priority list:
        # 1. Configured providers first (in order specified), then
        # 2. Any other providers with registered keys as implicit failover
        self.provider_priority = list(configured_with_keys)
        for p in ["gemini", "nova", "openai"]:
            if p not in self.provider_priority and self.key_pool.has_provider(p):
                self.provider_priority.append(p)

        # --- Client cache (key suffix -> SDK client) ---
        # We cache SDK clients to avoid re-creating httpx pools on every call.
        self._clients: dict = {}
        self._thread_local = threading.local()

        # Pre-initialize clients for all configured providers (warm start)
        self.client = None  # backward compat: first provider's client
        for provider in self.provider_priority:
            if provider == "openai" and self.key_pool.has_provider("openai"):
                import httpx
                import openai
                first_key = openai_keys[0] if openai_keys else ""
                client = openai.OpenAI(
                    api_key=first_key,
                    timeout=self.timeout,
                    max_retries=0,
                    http_client=httpx.Client(
                        limits=httpx.Limits(
                            max_connections=100,
                            max_keepalive_connections=20,
                            keepalive_expiry=30.0,
                        )
                    ),
                )
                self._clients[first_key[-6:]] = ("openai", client)
                if not self.client:
                    self.client = client
                logger.info(f"✓ OpenAI client initialized (model: {self.openai_model}, keys: {len(openai_keys)})")
            elif provider == "gemini" and self.key_pool.has_provider("gemini"):
                first_key = gemini_keys[0] if gemini_keys else ""
                genai.configure(api_key=first_key)
                client = genai.GenerativeModel(self.gemini_model)
                self._clients[first_key[-6:]] = ("gemini", client)
                if not self.client:
                    self.client = client
                logger.info(f"✓ Gemini client initialized (model: {self.gemini_model}, keys: {len(gemini_keys)})")
            elif provider == "nova" and self.key_pool.has_provider("nova"):
                import boto3
                client = boto3.client("bedrock-runtime", region_name=self.nova_region)
                self._clients["iam-cr"] = ("nova", client)
                if not self.client:
                    self.client = client
                logger.info(f"✓ Nova/Bedrock client initialized (model: {self.nova_model}, region: {self.nova_region})")

        # Log provider chain
        if len(self.provider_priority) > 1:
            logger.info(f"✓ AI provider chain: {' -> '.join(self.provider_priority)}")
        implicit_failover = [p for p in self.provider_priority if p not in self.ai_providers]
        if implicit_failover:
            logger.info(f"  (implicit failover: {', '.join(implicit_failover)})")

        # Check for Poppler availability
        self._check_poppler()

    def _check_poppler(self):
        """Check if Poppler is available for PDF processing"""
        try:
            import subprocess

            result = subprocess.run(["pdftoppm", "-v"], capture_output=True, text=True)
            logger.info("✓ Poppler is installed and available for PDF processing")
            return True
        except FileNotFoundError:
            logger.warning("⚠️  Poppler NOT found - PDF uploads will fail!")
            logger.warning("   Please install Poppler to enable PDF processing:")
            logger.warning(
                "   - Windows: Download from https://github.com/oschwartz10612/poppler-windows/releases/"
            )
            logger.warning("   - Set POPPLER_PATH in .env to the bin folder")
            logger.warning("   - Or add Poppler to your system PATH")
            return False
        except Exception as e:
            logger.warning(f"⚠️  Could not check Poppler: {e}")
            return False

    def _generate_cache_key(self, content: str, content_type: str = "image") -> str:
        """Generate cache key from content hash"""
        content_hash = hashlib.sha256(
            content.encode() if isinstance(content, str) else content
        ).hexdigest()
        return f"ai_extract:{content_type}:{content_hash[:16]}"

    def _get_cached_response(self, cache_key: str) -> Optional[dict]:
        """Get cached AI response if available"""
        if not self.cache:
            return None

        try:
            cached = self.cache.get(cache_key)
            if cached:
                logger.info(f"✓ Cache HIT: {cache_key[:30]}...")
                return json.loads(cached)
        except Exception as e:
            logger.warning(f"Cache read error: {e}")

        return None

    def _set_cached_response(self, cache_key: str, response: dict):
        """Cache AI response"""
        if not self.cache:
            return

        try:
            self.cache.setex(cache_key, self.cache_ttl, json.dumps(response))
            logger.debug(f"✓ Cached response: {cache_key[:30]}...")
        except Exception as e:
            logger.warning(f"Cache write error: {e}")

    def _call_with_retry(self, func, *args, **kwargs):
        """Call function with exponential backoff retry logic.
        Only retries on transient errors (429, 500, 502, 503, 504).
        Does NOT retry on client errors (400, 401, 403, 404) as they are deterministic.
        """
        last_exception = None

        for attempt in range(self.max_retries):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                last_exception = e
                # Don't retry client errors (400, 401, 403, 404) or SDK parameter errors
                error_str = str(e)
                is_client_error = any(
                    code in error_str
                    for code in ["400 -", "401 -", "403 -", "404 -",
                                 "status_code=400", "status_code=401",
                                 "invalid_request_error",
                                 "Missing required arguments"]
                ) or isinstance(e, TypeError)
                if is_client_error:
                    logger.error(f"❌ AI call failed with client error (no retry): {e}")
                    raise

                if attempt < self.max_retries - 1:
                    wait_time = self.retry_delay * (2**attempt)  # Exponential backoff
                    logger.warning(
                        f"⚠️  AI call failed (attempt {attempt + 1}/{self.max_retries}): {e}"
                    )
                    logger.info(f"   Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    logger.error(
                        f"❌ AI call failed after {self.max_retries} attempts: {e}"
                    )

        raise last_exception

    # ------------------------------------------------------------------
    # Thread-safe client management + multi-key failover
    # ------------------------------------------------------------------

    @property
    def _active_client(self):
        """Return the thread-local client, falling back to the default self.client"""
        return getattr(self._thread_local, "client", self.client)

    def _get_or_create_client(self, provider: str, api_key: str):
        """Get a cached SDK client or create a new one for the given key"""
        cache_key = api_key[-6:]
        if cache_key in self._clients:
            return self._clients[cache_key][1]

        if provider == "openai":
            import httpx
            import openai
            limits = httpx.Limits(
                max_connections=100,
                max_keepalive_connections=20,
                keepalive_expiry=30.0,
            )
            client = openai.OpenAI(
                api_key=api_key,
                timeout=self.timeout,
                max_retries=0,
                http_client=httpx.Client(limits=limits),
            )
        elif provider == "gemini":
            genai.configure(api_key=api_key)
            client = genai.GenerativeModel(self.gemini_model)
        elif provider == "nova":
            import boto3
            client = boto3.client("bedrock-runtime", region_name=self.nova_region)
        else:
            raise ValueError(f"Unsupported provider: {provider}")

        self._clients[cache_key] = (provider, client)
        return client

    def _get_extraction_method(self, provider: str, extraction_type: str):
        """Map (provider, extraction_type) to the correct method"""
        method_map = {
            ("gemini", "image"): self._extract_with_gemini,
            ("gemini", "pdf"): self._extract_pdf_with_gemini,
            ("gemini", "excel"): self._extract_excel_with_gemini,
            ("gemini", "columns"): self._ai_detect_columns_gemini,
            ("nova", "image"): self._extract_with_nova,
            ("nova", "pdf"): self._extract_pdf_with_nova,
            ("nova", "excel"): self._extract_excel_with_nova,
            ("nova", "columns"): self._ai_detect_columns_nova,
            ("openai", "image"): self._extract_with_openai,
            ("openai", "pdf"): self._extract_pdf_with_openai,
            ("openai", "excel"): self._extract_excel_with_openai,
            ("openai", "columns"): self._ai_detect_columns_openai,
        }
        return method_map.get((provider, extraction_type))

    def _is_rate_limit_error(self, error: Exception) -> bool:
        """Check if an exception is a 429 rate-limit error"""
        error_str = str(error)
        return "429" in error_str or "rate_limit" in error_str.lower() or "ResourceExhausted" in error_str or "ThrottlingException" in error_str

    def _is_client_error(self, error: Exception) -> bool:
        """Check if an exception is a 400 client error (bad request format, not a key issue)"""
        error_str = str(error)
        return ("400" in error_str and ("invalid" in error_str.lower() or "bad request" in error_str.lower())) or \
               "Invalid MIME type" in error_str or \
               "Missing required arguments" in error_str

    def _call_with_failover(self, extraction_type: str, *args, **kwargs):
        """
        Try all keys for the primary provider, then failover to secondary.

        For each key:
        1. Set the thread-local client
        2. Call the provider-specific extraction method (with retry)
        3. Report success/error to key pool

        Returns the extraction result (FinancialDocument or dict).
        Raises the last exception if all keys/providers are exhausted.
        """
        # Use full provider priority list (configured providers + implicit failover)
        # provider_priority already has configured providers first, then failover
        providers_to_try = list(self.provider_priority) if self.ai_failover_enabled else list(self.ai_providers)

        last_exception = None

        for provider in providers_to_try:
            method = self._get_extraction_method(provider, extraction_type)
            if not method:
                logger.warning(
                    f"No method for ({provider}, {extraction_type}), skipping"
                )
                continue

            # Try all keys for this provider
            tried_keys = 0
            while True:
                key_state = self.key_pool.get_next_key(provider)
                if not key_state:
                    break  # No more available keys for this provider
                tried_keys += 1

                # Set up thread-local client for this key
                client = self._get_or_create_client(provider, key_state.key)
                self._thread_local.client = client
                # Also set the model for the current provider
                self._thread_local.provider = provider

                try:
                    result = method(*args, **kwargs)
                    self.key_pool.report_success(key_state)
                    logger.info(
                        f"✓ AI call success: {provider} key {key_state.key_suffix} "
                        f"(type: {extraction_type})"
                    )
                    return result
                except Exception as e:
                    last_exception = e

                    # 400 client errors (bad format, unsupported MIME) are NOT key issues
                    # Don't punish the key and don't retry with different keys
                    if self._is_client_error(e):
                        logger.error(
                            f"❌ AI call failed with client error (no retry): {e}"
                        )
                        break  # Skip to next provider, same error with any key

                    is_rate_limit = self._is_rate_limit_error(e)
                    self.key_pool.report_error(key_state, is_rate_limit=is_rate_limit)
                    logger.warning(
                        f"⚠️  AI call failed: {provider} key {key_state.key_suffix} "
                        f"({'rate-limited' if is_rate_limit else type(e).__name__}): {e}"
                    )
                    # Continue to next key

            if tried_keys == 0 and provider == providers_to_try[0]:
                logger.warning(
                    f"No available keys for primary provider {provider}"
                )

            if provider != providers_to_try[-1]:
                logger.info(
                    f"Failing over from {provider} to next provider..."
                )

        # All providers exhausted
        if last_exception:
            raise last_exception
        raise ValueError("No AI API keys available for extraction")

    def process_document(self, file_path: str, user_company_info: dict = None, known_items: list = None) -> dict:
        """
        Process document and extract structured financial data
        Returns dict with status and extracted FinancialDocument
        Handles both local files and S3 files

        Args:
            file_path: Path to file (local or S3 key)
            user_company_info: Dict with company_name, legal_name, cnpj for income/expense detection
            known_items: List of dicts with name, alias, category, transaction_type, times_appeared
        """
        from config import settings
        from storage.s3_service import s3_storage
        import tempfile

        # Store user info and known items for use in extraction
        self.user_company_info = user_company_info
        self.known_items = known_items

        # Store original filename for context (income/expense detection from filename)
        self._original_filename = Path(file_path).name

        # If using S3, download file to temp location first
        temp_file = None
        if settings.use_s3:
            try:
                # Download from S3
                file_content = s3_storage.download_file(file_path)

                # Get file extension from S3 key
                extension = Path(file_path).suffix.lower()

                # Create temp file with same extension
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=extension)
                temp_file.write(file_content)
                temp_file.close()

                # Process the temp file
                file_path_obj = Path(temp_file.name)
            except Exception as e:
                if temp_file:
                    try:
                        Path(temp_file.name).unlink(missing_ok=True)
                    except:
                        pass
                raise Exception(f"Failed to download file from S3: {str(e)}")
        else:
            # Local filesystem
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path_obj.suffix.lower()

        try:
            # Process based on file type
            if extension in [".jpg", ".jpeg", ".png", ".webp"]:
                result = self._process_image(file_path_obj)
            elif extension == ".pdf":
                result = self._process_pdf(file_path_obj)
            elif extension in [".xlsx", ".xls"]:
                result = self._process_excel(file_path_obj)
            elif extension == ".xml":
                result = self._process_xml(file_path_obj)
            elif extension == ".ofx":
                result = self._process_ofx(file_path_obj)
            elif extension == ".ofc":
                result = self._process_ofc(file_path_obj)
            elif extension in [".doc", ".docx"]:
                result = self._process_docx(file_path_obj)
            elif extension == ".txt":
                result = self._process_txt(file_path_obj)
            else:
                raise ValueError(f"Unsupported file type: {extension}")

            return result

        finally:
            # Clean up temp file if using S3
            if temp_file:
                try:
                    Path(temp_file.name).unlink(missing_ok=True)
                except Exception as e:
                    logger.warning(f"Failed to delete temp file {temp_file.name}: {e}")

    def _process_image(self, image_path: Path) -> dict:
        """Process a single image file"""
        logger.info(f"📄 Processing image: {image_path.name}")

        try:
            with open(image_path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode("utf-8")

            ext = image_path.suffix.lower().replace(".", "")
            if ext == "jpg":
                ext = "jpeg"

            logger.info(f"🤖 Calling AI ({self.ai_provider}) for data extraction...")
            structured_data = self._extract_structured_data(image_data, ext)
            logger.info(f"✅ Image processing successful: {image_path.name}")

            return {
                "file_name": image_path.name,
                "file_type": "image",
                "status": "success",
                "extracted_data": structured_data,
            }
        except Exception as e:
            logger.error(
                f"❌ Error processing image {image_path.name}: {type(e).__name__}: {str(e)}"
            )
            import traceback

            logger.error(f"   Traceback:\n{traceback.format_exc()}")
            raise

    def _process_pdf_native(self, pdf_path: Path) -> dict:
        """
        Process PDF using native PDF support (OpenAI extracts text + images internally)

        Advantages over image conversion:
        - No Poppler dependency required
        - Smaller file sizes (PDF vs PNG images)
        - Faster processing
        - Better text extraction (native PDF text vs OCR)
        - Handles multi-page PDFs efficiently

        Supported models: GPT-4o, GPT-4o-mini, GPT-5, o1
        """
        logger.info(f"📄 Processing PDF natively (no image conversion): {pdf_path.name}")

        try:
            # Read PDF file and encode as base64
            with open(pdf_path, "rb") as f:
                pdf_data = base64.b64encode(f.read()).decode("utf-8")

            # Get file size for logging
            file_size_mb = len(pdf_data) * 3/4 / (1024 * 1024)  # base64 is ~4/3 of original size
            logger.info(f"📊 PDF size: {file_size_mb:.2f} MB")

            if file_size_mb > 20:
                logger.warning(f"⚠️ PDF is {file_size_mb:.2f} MB (>20MB recommended limit)")

            logger.info(f"🤖 Calling AI ({self.ai_provider}) for native PDF extraction...")
            structured_data = self._extract_structured_data_from_pdf(pdf_data)
            logger.info(f"✅ Native PDF processing successful: {pdf_path.name}")

            return {
                "file_name": pdf_path.name,
                "file_type": "pdf",
                "status": "success",
                "extracted_data": structured_data,
                "processing_method": "native_pdf",
            }
        except Exception as e:
            logger.error(
                f"❌ Error processing PDF natively {pdf_path.name}: {type(e).__name__}: {str(e)}"
            )
            raise  # Re-raise to trigger fallback to image conversion

    def _process_pdf(self, pdf_path: Path) -> dict:
        """
        Process PDF - tries native PDF support first for supported models,
        falls back to Poppler image conversion if native fails or model doesn't support PDFs
        """
        logger.info(f"📑 Processing PDF: {pdf_path.name}")

        # Try native PDF processing first - supported by GPT-4o, GPT-4o-mini, GPT-5, o1+ models
        # Only fine-tuned models don't support native PDF input
        if self.ai_provider == "openai":
            model_lower = self.openai_model.lower()
            supports_native_pdf = not any(tag in model_lower for tag in ["fine_tuned", "ft:"])
            if supports_native_pdf:
                try:
                    logger.info("📄 Attempting native PDF processing (no Poppler needed)...")
                    return self._process_pdf_native(pdf_path)
                except Exception as e:
                    logger.warning(f"⚠️ Native PDF failed: {e}")
                    logger.info("🔄 Falling back to Poppler image conversion...")
            else:
                logger.info(f"📄 Model {self.openai_model} does not support native PDF - using image conversion")

        # Fallback: Convert to images (original method)
        try:
            # On Windows, pdf2image needs explicit poppler_path
            poppler_path = os.getenv("POPPLER_PATH", None)

            logger.info("🔄 Converting PDF to images using Poppler...")
            if poppler_path:
                logger.info(f"   Using Poppler path from env: {poppler_path}")
                images = convert_from_path(str(pdf_path), poppler_path=poppler_path)
            else:
                logger.info("   Using Poppler from PATH")
                images = convert_from_path(str(pdf_path))
            logger.info(f"✓ PDF converted successfully: {len(images)} page(s)")

            # For multi-page PDFs, we'll process all pages and combine
            # For now, let's focus on single page or first page
            if len(images) > 1:
                logger.info(f"📚 Multi-page PDF detected ({len(images)} pages)")
                logger.info(f"🔄 Processing all pages...")

            all_structured_data = []

            for i, image in enumerate(images):
                logger.info(f"  📄 Processing page {i+1}/{len(images)}")

                # Resize image if too large (OpenAI has limits)
                max_dimension = 2000
                if image.width > max_dimension or image.height > max_dimension:
                    ratio = min(max_dimension / image.width, max_dimension / image.height)
                    new_size = (int(image.width * ratio), int(image.height * ratio))
                    image = image.resize(new_size, Image.Resampling.LANCZOS)
                    logger.info(f"  📐 Resized image to {new_size[0]}x{new_size[1]}")

                # Encode image directly in memory (no temp files = no race conditions)
                img_buffer = io.BytesIO()
                image.save(img_buffer, format="PNG", optimize=True)
                img_buffer.seek(0)
                image_data = base64.b64encode(img_buffer.read()).decode("utf-8")

                logger.info(f"  🤖 Calling AI ({self.ai_provider}) for page {i+1}...")
                page_data = self._extract_structured_data(
                    image_data, "png", page_num=i + 1
                )
                all_structured_data.append(page_data)
                logger.info(f"  ✅ Page {i+1} processed successfully")

            # For single page, return the data directly
            # For multi-page, return the first page's data (can be enhanced later)
            extracted_data = all_structured_data[0] if all_structured_data else None

            logger.info(f"✅ PDF processing successful (image conversion): {pdf_path.name}")
            return {
                "file_name": pdf_path.name,
                "file_type": "pdf",
                "page_count": len(images),
                "status": "success",
                "extracted_data": extracted_data,
                "all_pages_data": all_structured_data if len(images) > 1 else None,
                "processing_method": "image_conversion",
            }

        except Exception as e:
            logger.warning(
                f"⚠️ Poppler image conversion failed for {pdf_path.name}: {type(e).__name__}: {str(e)}"
            )
            logger.info("🔄 Falling back to text extraction (pypdf)...")

        # Third fallback: Extract text from PDF using pypdf and send as text
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(pdf_path))
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- Página {i+1} ---\n{page_text}")

            if text_parts:
                full_text = "\n\n".join(text_parts)
                logger.info(f"📝 Extracted {len(full_text)} chars of text from PDF ({len(reader.pages)} pages)")
                logger.info(f"🤖 Calling AI for text-based PDF extraction...")

                structured_data = self._extract_structured_data_from_excel(
                    f"Documento PDF: {pdf_path.name}\n\n{full_text}"
                )

                logger.info(f"✅ PDF processed via text extraction: {pdf_path.name}")
                return {
                    "file_name": pdf_path.name,
                    "file_type": "pdf",
                    "status": "success",
                    "extracted_data": structured_data,
                    "processing_method": "text_extraction",
                }
            else:
                logger.warning("⚠️ No extractable text found in PDF (scanned document?)")

        except Exception as text_err:
            logger.error(f"❌ Text extraction also failed: {type(text_err).__name__}: {str(text_err)}")

        # All methods failed
        return {
            "file_name": pdf_path.name,
            "file_type": "pdf",
            "status": "error",
            "error": "Não foi possível processar este PDF. Tente converter para imagem (JPG/PNG) antes de enviar.",
        }

    def _process_excel(self, excel_path: Path) -> dict:
        """Process Excel file (XLSX/XLS) — always sends through AI first"""
        logger.info(f"📊 Processing Excel file: {excel_path.name}")

        try:
            import pandas as pd

            # Read all sheets
            logger.info("🔄 Reading Excel file...")
            excel_file = pd.ExcelFile(excel_path)
            sheet_names = excel_file.sheet_names
            logger.info(f"✓ Found {len(sheet_names)} sheet(s): {sheet_names}")

            # Read ALL sheets and combine into one text for AI
            all_dfs = []
            for sheet_name in sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                if len(df) > 0:
                    all_dfs.append((sheet_name, df))

            if not all_dfs:
                return {
                    "file_name": excel_path.name,
                    "file_type": "excel",
                    "status": "error",
                    "error": "Arquivo Excel vazio — nenhuma aba contém dados",
                }

            # Check total row count to decide chunking strategy
            total_rows = sum(len(df) for _, df in all_dfs)
            CHUNK_THRESHOLD = 500  # rows above which we chunk to avoid blowing up AI context

            # PRIMARY PATH: Always send to AI for extraction
            logger.info(f"🤖 Sending Excel to AI ({self.ai_provider}) for extraction... ({total_rows} total rows)")
            try:
                if total_rows <= CHUNK_THRESHOLD:
                    # Small file — send everything in one shot
                    excel_text = self._all_sheets_to_text(all_dfs, excel_path.name)
                    structured_data = self._extract_structured_data_from_excel(excel_text)
                else:
                    # Large file — process in chunks and merge results
                    logger.info(f"📊 Large Excel ({total_rows} rows) — chunking into ~{CHUNK_THRESHOLD}-row batches")
                    structured_data = self._process_excel_chunked(all_dfs, excel_path.name)

                # Validate AI result has usable data
                ai_success = False
                if isinstance(structured_data, TransactionLedger):
                    # Chunked processing returns TransactionLedger directly
                    ai_success = structured_data.total_transactions > 0
                elif isinstance(structured_data, FinancialDocument):
                    has_txns = structured_data.transactions and len(structured_data.transactions) > 0
                    has_items = structured_data.line_items and len(structured_data.line_items) > 0
                    has_total = structured_data.total_amount and structured_data.total_amount > 0
                    ai_success = has_txns or has_items or has_total

                if ai_success:
                    logger.info(f"✅ AI extraction successful for {excel_path.name}")

                    # Override transaction types from filename/column context
                    # AI may not correctly detect income vs expense from Excel headers
                    original_name = getattr(self, '_original_filename', '') or ''
                    fn_lower = original_name.lower() if original_name else excel_path.name.lower()
                    primary_df = max(all_dfs, key=lambda x: len(x[1]))[1]
                    col_str = " ".join(str(c).lower() for c in primary_df.columns)
                    context_str = fn_lower + " " + col_str

                    _income_signals = ["recebimento", "receita", "faturamento", "total recebido", "valor recebido"]
                    _expense_signals = ["pagamento", "despesa", "compra", "total pago", "valor pago"]
                    detected_type = None
                    if any(s in context_str for s in _income_signals):
                        detected_type = "receita"
                    elif any(s in context_str for s in _expense_signals):
                        detected_type = "despesa"
                    # Also check cliente vs fornecedor in columns
                    if detected_type is None:
                        if "cliente" in col_str and "fornecedor" not in col_str:
                            detected_type = "receita"
                        elif "fornecedor" in col_str and "cliente" not in col_str:
                            detected_type = "despesa"

                    if detected_type and isinstance(structured_data, FinancialDocument):
                        # Override document-level type
                        if structured_data.transaction_type != detected_type:
                            logger.info(f"📋 Overriding AI transaction_type '{structured_data.transaction_type}' → '{detected_type}' from context")
                            structured_data.transaction_type = detected_type
                        # Override transaction-level types if they all have the wrong type
                        if structured_data.transactions:
                            for txn in structured_data.transactions:
                                if not txn.transaction_type or txn.transaction_type != detected_type:
                                    txn.transaction_type = detected_type

                    # Override document_type: Excel spreadsheets with multiple rows → transaction_ledger
                    if isinstance(structured_data, FinancialDocument) and structured_data.document_type == "statement":
                        if structured_data.transactions and len(structured_data.transactions) > 1:
                            logger.info(f"📋 Overriding document_type 'statement' → 'transaction_ledger' for Excel with {len(structured_data.transactions)} rows")
                            structured_data.document_type = "transaction_ledger"

                    return {
                        "file_name": excel_path.name,
                        "file_type": "excel",
                        "status": "success",
                        "extracted_data": structured_data,
                    }
                else:
                    logger.warning(f"⚠️ AI returned empty/unusable result, trying pandas fallback...")
            except Exception as e:
                logger.warning(f"⚠️ AI extraction failed ({e}), trying pandas fallback...")

            # FALLBACK: Use pandas-based ledger parsing
            # Use the first sheet with the most rows for ledger detection
            primary_df = max(all_dfs, key=lambda x: len(x[1]))[1]
            is_ledger = self._is_transaction_ledger(primary_df)

            if is_ledger:
                logger.info(f"📋 Fallback: pandas ledger parsing")
                structured_data = self._process_excel_as_ledger(primary_df, excel_path)

                # AI categorization for uncategorized transactions
                if isinstance(structured_data, TransactionLedger) and structured_data.transactions:
                    try:
                        structured_data.transactions = self._categorize_transactions_with_ai(
                            structured_data.transactions
                        )
                    except Exception as e:
                        logger.warning(f"⚠️ AI categorization skipped: {e}")

                # Check if fallback produced results
                ledger_failed = False
                if isinstance(structured_data, TransactionLedger):
                    if structured_data.total_transactions == 0:
                        ledger_failed = True
                    elif (structured_data.total_income == 0 and
                          structured_data.total_expense == 0):
                        ledger_failed = True

                if ledger_failed:
                    logger.error(f"❌ Both AI and pandas parsing failed for {excel_path.name}")
                    return {
                        "file_name": excel_path.name,
                        "file_type": "excel",
                        "status": "error",
                        "error": "Não foi possível extrair dados do arquivo Excel",
                    }
            else:
                logger.info(f"📄 Fallback: single document AI extraction")
                excel_text_single = self._dataframe_to_text(primary_df, excel_path.name)
                structured_data = self._extract_structured_data_from_excel(excel_text_single)

            logger.info(f"✅ Excel processing successful (fallback): {excel_path.name}")

            return {
                "file_name": excel_path.name,
                "file_type": "excel",
                "status": "success",
                "extracted_data": structured_data,
            }

        except ImportError:
            error_msg = (
                "pandas não está instalado. Execute: pip install pandas openpyxl"
            )
            logger.error(f"❌ {error_msg}")
            return {
                "file_name": excel_path.name,
                "file_type": "excel",
                "status": "error",
                "error": error_msg,
            }
        except Exception as e:
            logger.error(
                f"❌ Error processing Excel {excel_path.name}: {type(e).__name__}: {str(e)}"
            )
            import traceback

            logger.error(f"   Traceback:\n{traceback.format_exc()}")
            return {
                "file_name": excel_path.name,
                "file_type": "excel",
                "status": "error",
                "error": str(e),
            }

    def _dataframe_to_text(self, df, filename: str) -> str:
        """Convert pandas DataFrame to text representation for AI"""
        text_parts = [f"Arquivo Excel: {filename}\n"]
        text_parts.append(f"Total de linhas: {len(df)}\n")
        text_parts.append(f"Colunas: {', '.join(df.columns.astype(str))}\n\n")

        # Provide statistical summary for large datasets
        text_parts.append("Resumo Estatístico:\n")
        text_parts.append(df.describe(include="all").to_string())
        text_parts.append("\n\n")

        text_parts.append("Dados:\n")

        # Send ALL rows to AI - no sampling, clients need complete extraction
        text_parts.append(df.to_string(index=False))

        return "\n".join(text_parts)

    def _all_sheets_to_text(self, all_dfs: list, filename: str) -> str:
        """Convert all Excel sheets to text representation for AI"""
        text_parts = [f"Arquivo Excel: {filename}\n"]
        text_parts.append(f"Total de abas: {len(all_dfs)}\n\n")

        for sheet_name, df in all_dfs:
            text_parts.append(f"{'='*60}\n")
            text_parts.append(f"ABA: {sheet_name}\n")
            text_parts.append(f"Linhas: {len(df)}, Colunas: {len(df.columns)}\n")
            text_parts.append(f"Colunas: {', '.join(df.columns.astype(str))}\n\n")

            # Send ALL rows - no sampling
            text_parts.append(df.to_string(index=False))

            text_parts.append("\n\n")

        return "\n".join(text_parts)

    def _process_excel_chunked(self, all_dfs: list, filename: str):
        """
        Process a large Excel file by splitting into chunks, extracting each
        chunk with AI, and merging the results into a single TransactionLedger.

        This avoids blowing up AI context windows on 1000+ row spreadsheets
        while still extracting every single row.
        """
        import pandas as pd
        from models import TransactionLedger, Transaction, DateRangeSummary

        CHUNK_SIZE = 500
        all_transactions = []
        chunk_index = 0

        for sheet_name, df in all_dfs:
            total_rows = len(df)
            if total_rows == 0:
                continue

            # Split sheet into chunks
            for start in range(0, total_rows, CHUNK_SIZE):
                chunk_df = df.iloc[start:start + CHUNK_SIZE]
                chunk_index += 1
                end = min(start + CHUNK_SIZE, total_rows)
                logger.info(
                    f"  Chunk {chunk_index}: sheet '{sheet_name}' rows {start+1}-{end} of {total_rows}"
                )

                # Build text for this chunk (include column headers for context)
                text_parts = [
                    f"Arquivo Excel: {filename}\n",
                    f"ABA: {sheet_name} (linhas {start+1} a {end} de {total_rows})\n",
                    f"Colunas: {', '.join(df.columns.astype(str))}\n\n",
                    "Dados:\n",
                    chunk_df.to_string(index=False),
                ]
                chunk_text = "\n".join(text_parts)

                try:
                    result = self._extract_structured_data_from_excel(chunk_text)

                    # Collect transactions from the AI result
                    if isinstance(result, FinancialDocument) and result.transactions:
                        all_transactions.extend(result.transactions)
                    elif hasattr(result, "transactions") and result.transactions:
                        all_transactions.extend(result.transactions)
                except Exception as e:
                    logger.warning(f"  Chunk {chunk_index} AI extraction failed: {e}")
                    # Continue with remaining chunks — partial extraction is better than none

        if not all_transactions:
            # All chunks failed — return empty FinancialDocument so caller hits fallback
            return FinancialDocument(
                document_type="transaction_ledger",
                document_number=filename,
            )

        # Merge all chunk transactions into a single TransactionLedger
        total_income = sum(
            t.amount for t in all_transactions
            if t.transaction_type in ("income", "receita")
        )
        total_expense = sum(
            t.amount for t in all_transactions
            if t.transaction_type in ("expense", "despesa")
        )

        dates = [t.date for t in all_transactions if t.date]
        date_range = DateRangeSummary()
        if dates:
            date_range.start_date = min(dates)
            date_range.end_date = max(dates)

        ledger = TransactionLedger(
            document_type="transaction_ledger",
            document_number=filename,
            total_transactions=len(all_transactions),
            total_income=total_income,
            total_expense=total_expense,
            net_balance=total_income - total_expense,
            transactions=all_transactions,
            date_range=date_range,
        )

        logger.info(
            f"Chunked extraction complete: {len(all_transactions)} transactions "
            f"from {chunk_index} chunks (income={total_income}, expense={total_expense})"
        )
        return ledger

    def _is_transaction_ledger(self, df) -> bool:
        """
        Detect if DataFrame is a transaction ledger (multiple transaction rows)
        vs a single document.  Checks BOTH column headers AND actual row data,
        since headers may be on a non-zero row that pandas didn't auto-detect.
        """
        import pandas as pd

        df_clean = df.dropna(how="all")

        # Very small files (< 3 data rows) are never ledgers
        if len(df_clean) < 3:
            return False

        ledger_keywords = [
            "data", "date", "dt",
            "valor", "amount", "value", "montante",
            "total", "parcela",
            "descri", "historico", "histórico", "memo",
            "category", "categoria", "natureza", "classificação",
            "debito", "débito", "credito", "crédito",
            "debit", "credit",
            "saldo", "balance",
            "pago", "recebido",
            "fornecedor", "cliente", "pagador",
            "banco", "bank", "conta", "account",
            "vencimento", "competência", "competencia",
            "tipo", "type",
        ]

        # Strategy 1: Check pandas-detected column names (row 0 as header)
        column_str = " ".join(str(col).lower() for col in df.columns)
        keyword_matches = sum(1 for kw in ledger_keywords if kw in column_str)

        if keyword_matches >= 2:
            logger.info(f"📋 Ledger detection: {keyword_matches} keywords in column names")
            return True

        # Strategy 2: Scan first 15 rows for a header row (real header may not be row 0)
        # This catches spreadsheets with title rows, logos, blank rows before the data
        for i in range(min(15, len(df))):
            row = df.iloc[i]
            row_str = " ".join(str(val).lower() for val in row if pd.notna(val))
            row_matches = sum(1 for kw in ledger_keywords if kw in row_str)
            if row_matches >= 2:
                logger.info(f"📋 Ledger detection: {row_matches} keywords in row {i}")
                return True

        # Strategy 3: Structural heuristics - many rows + numeric column = ledger
        numeric_cols = df.select_dtypes(include=["float64", "int64"]).columns
        if len(numeric_cols) >= 1 and len(df_clean) > 20:
            logger.info(
                f"📋 Ledger detection: {len(numeric_cols)} numeric cols, {len(df_clean)} rows"
            )
            return True

        # Strategy 4: Check if many cells in a column contain monetary-looking strings
        # (catches cases where pandas reads amounts as strings like "1.234,56")
        for col in df.columns:
            sample = df[col].dropna().head(20).astype(str)
            import re
            monetary_pattern = re.compile(r'^-?\d{1,3}(\.\d{3})*(,\d{2})?$|^-?R?\$?\s?\d')
            monetary_count = sum(1 for val in sample if monetary_pattern.match(val.strip()))
            if monetary_count >= 3 and len(df_clean) >= 5:
                logger.info(f"📋 Ledger detection: {monetary_count} monetary values in column '{col}'")
                return True

        return False

    def _clean_dataframe(self, df) -> any:
        """
        Aggressively clean dataframe to remove garbage rows
        Handles: empty rows, whitespace-only, summary rows, formula errors
        """
        import pandas as pd

        # Remove completely empty rows
        df = df.dropna(how="all")

        # Remove rows where all values are just whitespace
        def is_whitespace_row(row):
            for val in row:
                if pd.notna(val):
                    val_str = str(val).strip()
                    if val_str and val_str not in ["", " ", "nan", "NaN", "None"]:
                        return False
            return True

        df = df[~df.apply(is_whitespace_row, axis=1)]

        # Remove rows with Excel error values
        def has_excel_error(row):
            for val in row:
                if pd.notna(val):
                    val_str = str(val).upper()
                    if any(err in val_str for err in ["#N/A", "#DIV/0", "#VALUE", "#REF", "#NAME", "#NUM", "#NULL"]):
                        return True
            return False

        df = df[~df.apply(has_excel_error, axis=1)]

        # Remove obvious summary rows (TOTAL, SUBTOTAL, etc.)
        def is_summary_row(row):
            for val in row:
                if pd.notna(val) and isinstance(val, str):
                    val_lower = val.lower().strip()
                    if val_lower in ["total", "subtotal", "soma", "sum", "grand total", "totals", "saldo"]:
                        return True
            return False

        df = df[~df.apply(is_summary_row, axis=1)]

        return df.reset_index(drop=True)

    def _process_excel_as_ledger(self, df, excel_path: Path) -> TransactionLedger:
        """
        Process Excel file as a transaction ledger
        Parses each row as a transaction and creates aggregated summaries
        """
        import pandas as pd

        # Clean up the dataframe
        df_clean = df.dropna(how="all")  # Remove empty rows

        # Extended header keywords for Brazilian accounting exports
        header_keywords = [
            "data", "date", "dt",
            "valor", "value", "montante", "amount",
            "descri", "description", "historico", "histórico", "memo",
            "categoria", "category", "natureza", "classificação", "classificacao",
            "banco", "bank", "conta", "account",
            "comp", "competência", "competencia",
            "débito", "debito", "crédito", "credito",
            "debit", "credit",
            "saldo", "balance",
            "tipo", "type",
            "vencimento", "pagamento", "recebimento",
            "fornecedor", "cliente", "pagador",
            "parcela", "total",
            "centro de custo", "plano de contas",
        ]

        # Robust header detection - try multiple strategies
        header_row = None

        # Strategy 1: Check if pandas already got it right (row 0 as header)
        # If column names already contain keywords, header_row stays None (pandas got it)
        col_str = " ".join(str(col).lower() for col in df.columns)
        pandas_keyword_hits = sum(1 for kw in header_keywords if kw in col_str)
        # Also check if columns look auto-generated (Unnamed: 0, Column1, etc.)
        auto_generated = sum(1 for col in df.columns if str(col).startswith(("Unnamed", "Column")))

        if pandas_keyword_hits >= 2 and auto_generated < len(df.columns) * 0.5:
            # Pandas detected headers correctly, no need to re-read
            logger.info(f"📋 Pandas auto-detected header with {pandas_keyword_hits} keywords")
        else:
            # Strategy 2: Scan rows 0-20 for the real header row
            for i in range(min(20, len(df))):
                row = df.iloc[i]
                row_values = [str(val).lower() for val in row if pd.notna(val)]
                row_str = " ".join(row_values)

                # Count keyword matches
                keyword_matches = sum(1 for kw in header_keywords if kw in row_str)

                # Also count how many cells look like text (not numbers)
                text_count = sum(
                    1 for val in row
                    if pd.notna(val) and isinstance(val, str) and len(val.strip()) > 1
                )

                # Accept header if: has keywords, OR has many text cells that aren't data
                if keyword_matches >= 2:
                    header_row = i
                    logger.info(f"📋 Detected header row at index {i} ({keyword_matches} keyword matches)")
                    break
                elif keyword_matches >= 1 and text_count >= 2:
                    header_row = i
                    logger.info(f"📋 Detected header row at index {i} ({keyword_matches} keywords + {text_count} text cells)")
                    break

        # Re-read with correct header
        if header_row is not None and header_row > 0:
            # Read without header first
            df_no_header = pd.read_excel(excel_path, header=None)
            # Get column names from detected header row
            column_names = df_no_header.iloc[header_row].tolist()
            # Clean up column names - convert to strings and filter out numeric-only names
            cleaned_columns = []
            for i, col in enumerate(column_names):
                if isinstance(col, (int, float)) and not pd.isna(col):
                    # Numeric column name - skip or use generic name
                    cleaned_columns.append(f"Column_{i}")
                elif pd.isna(col):
                    cleaned_columns.append(f"Unnamed_{i}")
                else:
                    cleaned_columns.append(str(col))

            # Get data starting from row after header
            df_clean = df_no_header.iloc[header_row + 1:].copy()
            df_clean.columns = cleaned_columns
            df_clean = df_clean.reset_index(drop=True)
            df_clean = self._clean_dataframe(df_clean)
            logger.info(f"✓ Re-read Excel with header row {header_row}, columns: {list(df_clean.columns)[:5]}")
        else:
            df_clean = self._clean_dataframe(df_clean)

        # Auto-detect columns
        date_col = None
        description_col = None
        amount_col = None
        category_col = None
        counterparty_col = None

        # --- Date column: prefer accounting date over due date ---
        date_priority_keywords = ["competência", "competencia", "pagamento", "recebimento", "lançamento", "lancamento"]
        date_fallback_keywords = ["data", "date", "dt"]
        for priority_kw in date_priority_keywords:
            for col in df_clean.columns:
                if priority_kw in str(col).lower():
                    date_col = col
                    break
            if date_col:
                break
        if date_col is None:
            for col in df_clean.columns:
                col_str = str(col).lower()
                if any(kw in col_str for kw in date_fallback_keywords):
                    date_col = col
                    break

        # --- Description column ---
        for col in df_clean.columns:
            col_str = str(col).lower()
            if any(
                keyword in col_str
                for keyword in ["descri", "description", "historico", "histórico", "memo", "observ"]
            ):
                description_col = col
                break

        # --- Amount column: priority-ranked detection ---
        # Also detect debit/credit dual-column pattern (common in bank exports)
        debit_col = None
        credit_col = None

        # Priority 1: "total recebido", "total pago", "valor total pago" (the actual final amount)
        # Priority 2: "valor pago", "valor recebido"
        # Priority 3: generic "valor", "amount", "montante", etc.
        # Priority 4: "total da parcela", "total" (installment, less ideal)
        amount_priority_tiers = [
            ["total recebido", "total pago", "valor total pago", "valor total recebido"],
            ["valor pago", "valor recebido"],
            ["valor", "value", "amount", "montante"],
            ["total da parcela", "total", "parcela"],
        ]
        for tier in amount_priority_tiers:
            for kw in tier:
                for col in df_clean.columns:
                    if kw in str(col).lower():
                        amount_col = col
                        break
                if amount_col:
                    break
            if amount_col:
                break

        # If no single amount column, check for debit/credit dual columns
        if amount_col is None:
            for col in df_clean.columns:
                col_lower = str(col).lower()
                if any(kw in col_lower for kw in ["débito", "debito", "debit"]):
                    debit_col = col
                elif any(kw in col_lower for kw in ["crédito", "credito", "credit"]):
                    credit_col = col
            if debit_col and credit_col:
                logger.info(f"📋 Detected debit/credit dual columns: {debit_col} / {credit_col}")
                # We'll handle this in the row loop below

        # Fallback: pick numeric column with most non-null values (skip mostly-NaN cols)
        if amount_col is None and not (debit_col and credit_col):
            numeric_cols = df_clean.select_dtypes(include=["float64", "int64"]).columns
            best_col, best_count = None, 0
            for col in numeric_cols:
                count = int(df_clean[col].notna().sum())
                if count > best_count:
                    best_count = count
                    best_col = col
            if best_col and best_count > 0:
                amount_col = best_col

        # Last-resort fallback: check if any column has string values that look like money
        # (pandas may have read amounts as strings if formatting is unusual)
        if amount_col is None and not (debit_col and credit_col):
            import re
            monetary_re = re.compile(r'^-?\(?\d{1,3}(\.\d{3})*(,\d{2})?\)?$|^-?R?\$?\s?\d')
            for col in df_clean.columns:
                if col in [date_col, description_col, category_col, counterparty_col]:
                    continue
                sample = df_clean[col].dropna().head(10).astype(str)
                hits = sum(1 for v in sample if monetary_re.match(v.strip()))
                if hits >= 3:
                    amount_col = col
                    logger.info(f"📋 Detected string-format monetary column: '{col}' ({hits}/10 samples match)")
                    break

        # --- Category column: broader keywords ---
        category_keywords = [
            "categoria", "categ", "natureza", "plano de contas",
            "grupo", "centro de custo", "classificação", "classificacao",
            "tipo", "type", "class",
        ]
        for col in df_clean.columns:
            col_str = str(col).lower()
            if any(keyword in col_str for keyword in category_keywords):
                category_col = col
                break

        # --- Counterparty column (supplier/client) ---
        for col in df_clean.columns:
            col_str = str(col).lower()
            if any(kw in col_str for kw in ["fornecedor", "cliente", "pagador", "supplier", "customer"]):
                counterparty_col = col
                break

        # --- Explicit "Tipo" column (from our template or user-provided) ---
        type_col = None
        for col in df_clean.columns:
            col_str = str(col).lower().strip()
            if col_str == "tipo" or col_str == "type" or col_str == "transaction_type":
                type_col = col
                break

        # --- Reference/document number column ---
        reference_col = None
        for col in df_clean.columns:
            col_str = str(col).lower()
            if any(kw in col_str for kw in ["nº documento", "n documento", "numero documento", "document", "referência", "referencia", "reference", "nf"]):
                reference_col = col
                break

        # --- Transaction type from filename/column context ---
        force_type = None
        # Use original filename (preserved from S3 key or upload path) for better detection
        original_name = getattr(self, '_original_filename', None) or ''
        filename_lower = original_name.lower() if original_name else (
            excel_path.name.lower() if hasattr(excel_path, 'name') else str(excel_path).lower()
        )
        # Also check sheet name for context
        sheet_name_lower = str(getattr(df, 'name', '') or '').lower() if hasattr(df, 'name') else ''
        filename_context = filename_lower + ' ' + sheet_name_lower
        if any(kw in filename_context for kw in ["recebimento", "receita", "faturamento", "income", "receivable", "recebivel"]):
            force_type = "receita"
        elif any(kw in filename_context for kw in ["pagamento", "despesa", "compra", "expense", "payable"]):
            force_type = "despesa"
        # Also detect from column names if filename didn't match
        if force_type is None:
            column_str_lower = " ".join(str(c).lower() for c in df_clean.columns)
            # Strong income signals: "recebido", "total recebido", "cliente" without "fornecedor"
            income_col_signals = [
                "total recebido", "valor recebido", "recebido",
                "receita", "faturamento", "recebimento",
            ]
            expense_col_signals = [
                "total pago", "valor pago", "pagamento",
                "despesa", "compra",
            ]
            if any(sig in column_str_lower for sig in income_col_signals):
                force_type = "receita"
            elif any(sig in column_str_lower for sig in expense_col_signals):
                force_type = "despesa"
            elif "cliente" in column_str_lower and "fornecedor" not in column_str_lower:
                force_type = "receita"
            elif "fornecedor" in column_str_lower and "cliente" not in column_str_lower:
                force_type = "despesa"

        if force_type:
            logger.info(f"📋 Transaction type from context: {force_type}")

        logger.info(
            f"📋 Detected columns - Date: {date_col}, Desc: {description_col}, "
            f"Amount: {amount_col}, Debit: {debit_col}, Credit: {credit_col}, "
            f"Category: {category_col}, Counterparty: {counterparty_col}"
        )

        # Determine if we have usable amount data
        has_amount = amount_col is not None or (debit_col is not None and credit_col is not None)

        # If we're missing critical columns, use AI to detect them
        if not date_col or not has_amount:
            logger.info("🤖 Using AI to detect missing columns...")
            try:
                ai_detected_cols = self._ai_detect_columns(df_clean)
                if not date_col and ai_detected_cols.get('date'):
                    date_col = ai_detected_cols['date']
                    logger.info(f"✓ AI detected date column: {date_col}")
                if not has_amount and ai_detected_cols.get('amount'):
                    amount_col = ai_detected_cols['amount']
                    has_amount = True
                    logger.info(f"✓ AI detected amount column: {amount_col}")
                if not description_col and ai_detected_cols.get('description'):
                    description_col = ai_detected_cols['description']
                if not category_col and ai_detected_cols.get('category'):
                    category_col = ai_detected_cols['category']
            except Exception as e:
                logger.error(f"❌ AI column detection failed: {e}")

        # Last resort: if still no amount column, pick numeric col with most data
        if not has_amount:
            numeric_cols = df_clean.select_dtypes(include=["float64", "int64"]).columns
            best_col, best_count = None, 0
            for col in numeric_cols:
                count = int(df_clean[col].notna().sum())
                if count > best_count:
                    best_count = count
                    best_col = col
            if best_col and best_count > 0:
                amount_col = best_col
                has_amount = True
                logger.warning(f"⚠️ Using numeric column with most data as amount: {amount_col} ({best_count} values)")

        # Parse transactions
        transactions = []
        dates = []
        skipped_reasons = {"no_amount_col": 0, "amount_parse_fail": 0, "amount_zero": 0, "exception": 0}
        use_debit_credit = (amount_col is None and debit_col is not None and credit_col is not None)

        for idx, row in df_clean.iterrows():
            try:
                # Extract amount - be defensive about column access
                amount = None
                debit_credit_type = None  # Track if amount came from debit or credit col

                if use_debit_credit:
                    # Dual debit/credit columns: use whichever is non-empty
                    try:
                        dval = row[debit_col]
                        cval = row[credit_col]
                        if pd.notna(dval):
                            parsed = self._parse_cell_as_decimal(dval)
                            if parsed and parsed != 0:
                                amount = -abs(parsed)  # Debits are expenses (negative)
                                debit_credit_type = "despesa"
                        if amount is None and pd.notna(cval):
                            parsed = self._parse_cell_as_decimal(cval)
                            if parsed and parsed != 0:
                                amount = abs(parsed)  # Credits are income (positive)
                                debit_credit_type = "receita"
                    except Exception:
                        skipped_reasons["amount_parse_fail"] += 1
                        continue
                elif amount_col is not None:
                    try:
                        val = row[amount_col]
                        if pd.notna(val):
                            if isinstance(val, pd.Series):
                                val = val.iloc[0] if len(val) > 0 else None
                            if val is not None:
                                amount = self._parse_cell_as_decimal(val)
                    except Exception as e:
                        skipped_reasons["amount_parse_fail"] += 1
                        continue
                else:
                    skipped_reasons["no_amount_col"] += 1
                    continue

                if amount is None or amount == 0:
                    skipped_reasons["amount_zero"] += 1
                    continue

                # Extract date - defensive
                date_str = None
                if date_col is not None:
                    try:
                        val = row[date_col]
                        if isinstance(val, pd.Series):
                            val = val.iloc[0] if len(val) > 0 else None
                        if pd.notna(val):
                            date_obj = pd.to_datetime(val, dayfirst=True)
                            date_str = date_obj.strftime("%Y-%m-%d")
                            dates.append(date_obj)
                    except:
                        pass

                # Extract description - defensive
                description = None
                if description_col is not None:
                    try:
                        val = row[description_col]
                        if isinstance(val, pd.Series):
                            val = val.iloc[0] if len(val) > 0 else None
                        if pd.notna(val):
                            description = str(val).strip()[:200]
                    except:
                        pass

                # If no description column detected, concatenate all text cells as description
                if description is None:
                    text_parts = []
                    for col in df_clean.columns:
                        if col not in [date_col, amount_col, category_col, counterparty_col]:
                            try:
                                val = row[col]
                                if pd.notna(val) and isinstance(val, str) and val.strip():
                                    text_parts.append(val.strip())
                            except:
                                pass
                    if text_parts:
                        description = " | ".join(text_parts)[:200]

                # Extract category - defensive
                category = None
                if category_col is not None:
                    try:
                        val = row[category_col]
                        if isinstance(val, pd.Series):
                            val = val.iloc[0] if len(val) > 0 else None
                        if pd.notna(val):
                            category = str(val).strip()[:100]
                    except:
                        pass

                # Extract counterparty (supplier/client) into notes
                counterparty = None
                if counterparty_col is not None:
                    try:
                        val = row[counterparty_col]
                        if isinstance(val, pd.Series):
                            val = val.iloc[0] if len(val) > 0 else None
                        if pd.notna(val):
                            counterparty = str(val).strip()[:200]
                    except:
                        pass

                # Read explicit type from "Tipo" column if present
                row_type = None
                if type_col is not None:
                    try:
                        val = row[type_col]
                        if isinstance(val, pd.Series):
                            val = val.iloc[0] if len(val) > 0 else None
                        if pd.notna(val):
                            row_type_str = str(val).lower().strip()
                            type_map = {
                                "receita": "receita", "income": "receita",
                                "despesa": "despesa", "expense": "despesa", "gasto": "despesa",
                                "custo": "custo", "cost": "custo",
                                "investimento": "investimento",
                                "perda": "perda",
                            }
                            row_type = type_map.get(row_type_str)
                    except:
                        pass

                # Read reference/document number if present
                reference = None
                if reference_col is not None:
                    try:
                        val = row[reference_col]
                        if isinstance(val, pd.Series):
                            val = val.iloc[0] if len(val) > 0 else None
                        if pd.notna(val):
                            reference = str(val).strip()[:100]
                    except:
                        pass

                # Determine transaction type (priority: row > force_type > debit/credit > sign)
                if row_type:
                    transaction_type = row_type
                    amount = abs(amount)
                elif force_type:
                    transaction_type = force_type
                    amount = abs(amount)
                elif debit_credit_type:
                    transaction_type = debit_credit_type
                    amount = abs(amount)
                else:
                    transaction_type = "despesa" if amount < 0 else "receita"
                    amount = abs(amount)

                transaction = Transaction(
                    date=date_str,
                    description=description or f"Linha {idx + 1}",
                    category=category or "nao_categorizado",
                    amount=amount,
                    transaction_type=transaction_type,
                    counterparty=counterparty,
                    reference=reference,
                )

                transactions.append(transaction)

            except Exception as e:
                skipped_reasons["exception"] += 1
                continue

        # Diagnostic logging - always show skip reasons
        total_skipped = sum(skipped_reasons.values())
        if total_skipped > 0:
            logger.info(
                f"📊 Row parsing stats: {len(transactions)} parsed, {total_skipped} skipped "
                f"(no_amount_col={skipped_reasons['no_amount_col']}, "
                f"parse_fail={skipped_reasons['amount_parse_fail']}, "
                f"zero={skipped_reasons['amount_zero']}, "
                f"exception={skipped_reasons['exception']})"
            )

        if len(transactions) == 0:
            logger.error(f"❌ Parsed 0 transactions from {len(df_clean)} rows - column detection likely failed")
            logger.error(f"   Detected columns: date={date_col}, amount={amount_col}, desc={description_col}, cat={category_col}")
            # Log sample of what the data actually looks like
            if len(df_clean) > 0:
                sample_row = df_clean.iloc[0]
                logger.error(f"   Sample row 0: {dict(sample_row)}")
        elif len(transactions) < len(df_clean) * 0.1:
            logger.warning(
                f"⚠️ Only parsed {len(transactions)}/{len(df_clean)} rows ({len(transactions)*100//len(df_clean)}%) - "
                f"possible format issue"
            )

        # Calculate summaries
        total_income = Decimal("0")
        total_expense = Decimal("0")
        category_totals = defaultdict(
            lambda: {"income": Decimal("0"), "expense": Decimal("0"), "count": 0}
        )

        for t in transactions:
            # Normalize Portuguese transaction types to English for aggregation
            is_income = t.transaction_type in ("income", "receita")
            if is_income:
                total_income += t.amount
            else:
                total_expense += t.amount

            agg_key = "income" if is_income else "expense"
            category_totals[t.category][agg_key] += t.amount
            category_totals[t.category]["count"] += 1

        # Create category summaries
        by_category = []
        for category, data in category_totals.items():
            # Determine dominant transaction type for this category
            if data["income"] > data["expense"]:
                trans_type = "income"
                amount = data["income"]
            else:
                trans_type = "expense"
                amount = data["expense"]

            by_category.append(
                CategorySummary(
                    category=category,
                    total_amount=amount,
                    count=data["count"],
                    transaction_type=trans_type,
                )
            )

        # Sort by amount descending
        by_category.sort(key=lambda x: x.total_amount, reverse=True)

        # Date range
        date_range = DateRangeSummary()
        if dates:
            date_range.start_date = min(dates).strftime("%Y-%m-%d")
            date_range.end_date = max(dates).strftime("%Y-%m-%d")
            date_range.total_days = (max(dates) - min(dates)).days

        # Create ledger
        ledger = TransactionLedger(
            file_name=excel_path.name,
            total_transactions=len(transactions),
            date_range=date_range,
            total_income=total_income,
            total_expense=total_expense,
            net_balance=total_income - total_expense,
            by_category=by_category,
            transactions=transactions,
        )

        logger.info(
            f"✅ Ledger summary - Income: {total_income}, Expense: {total_expense}, Balance: {total_income - total_expense}"
        )

        # Clean up dataframe references to release file handle
        del df_clean
        if 'df_no_header' in locals():
            del df_no_header
        import gc
        gc.collect()

        return ledger

    def _categorize_transactions_with_ai(self, transactions: list) -> list:
        """
        Batch-categorize transactions using AI.
        Groups unique uncategorized descriptions and asks AI to assign V2 category names.
        Called after Excel ledger parsing when transactions lack proper categories.
        """
        UNCATEGORIZED = {"nao_categorizado", "uncategorized", ""}
        uncategorized_descriptions = sorted({
            t.description for t in transactions
            if (not t.category or t.category.strip().lower() in UNCATEGORIZED) and t.description
        })

        if not uncategorized_descriptions:
            return transactions

        logger.info(f"🤖 AI categorizing {len(uncategorized_descriptions)} unique descriptions...")

        CATEGORIES_TEXT = (
            "RECEITA: receita_vendas_produtos, receita_servicos, receita_locacao, receita_comissoes, receita_contratos_recorrentes\n"
            "DEDUÇÕES: impostos_sobre_vendas, devolucoes, descontos_concedidos\n"
            "CUSTOS VARIÁVEIS: cmv, csp, materia_prima, insumos, comissoes_sobre_vendas\n"
            "CUSTOS FIXOS PRODUÇÃO: salarios_producao, encargos_sociais_producao, energia_producao, manutencao_equipamentos_producao\n"
            "DESPESAS ADMIN: salarios_administrativos, pro_labore, encargos_sociais_administrativos, aluguel, condominio, agua_energia, material_escritorio, honorarios_contabeis, sistemas_softwares, telefonia_internet\n"
            "DESPESAS COMERCIAIS: marketing_publicidade, propaganda_digital, comissao_vendas, fretes, representantes_comerciais\n"
            "FINANCEIRO: receita_financeira, juros_ativos, descontos_obtidos, juros_passivos, tarifas_bancarias, iof, multas_encargos\n"
            "TRIBUTOS: irpj, csll, simples_nacional, iptu, taxas_municipais\n"
            "OUTRAS: recuperacao_despesas, venda_imobilizado, indenizacoes_recebidas, outras_receitas_eventuais, perdas, indenizacoes_pagas, doacoes, provisoes, depreciacao, amortizacao, outras_despesas_operacionais\n"
            "FALLBACK: nao_categorizado (ONLY when none above fit)"
        )

        BATCH_SIZE = 80
        description_to_category: dict = {}

        for batch_start in range(0, len(uncategorized_descriptions), BATCH_SIZE):
            batch = uncategorized_descriptions[batch_start: batch_start + BATCH_SIZE]
            descriptions_json = json.dumps(batch, ensure_ascii=False)

            prompt = (
                "You are a Brazilian accounting assistant. Categorize each transaction description below "
                "using EXACTLY one category key from the V2 Plano de Contas list.\n\n"
                f"V2 categories:\n{CATEGORIES_TEXT}\n\n"
                "Return ONLY a valid JSON object mapping each description exactly as given to its category key. "
                "No markdown, no explanation.\n"
                f"Descriptions: {descriptions_json}"
            )

            try:
                if self.ai_provider == "openai":
                    response = self._active_client.chat.completions.create(
                        model=self.openai_model,
                        messages=[{"role": "user", "content": prompt}],
                        max_completion_tokens=4000,
                        store=False,
                    )
                    raw = response.choices[0].message.content or "{}"
                elif self.ai_provider == "gemini":
                    response = self._active_client.generate_content(
                        prompt,
                        generation_config=genai.types.GenerationConfig(
                            max_output_tokens=4000,
                            temperature=0.1,
                        ),
                    )
                    raw = response.text or "{}"
                elif self.ai_provider == "nova":
                    response = self._active_client.converse(
                        modelId=self.nova_model,
                        messages=[{
                            "role": "user",
                            "content": [{"text": prompt}],
                        }],
                        inferenceConfig={"maxTokens": 4000, "temperature": 0.1},
                    )
                    raw = response["output"]["message"]["content"][0]["text"] if response.get("output") else "{}"
                else:
                    raise ValueError(f"Unsupported AI provider: {self.ai_provider}")

                # Strip markdown fences if present
                raw = raw.strip()
                if raw.startswith("```"):
                    raw = re.sub(r"^```[a-z]*\n?", "", raw)
                    raw = re.sub(r"\n?```$", "", raw)

                batch_result = json.loads(raw)
                description_to_category.update(batch_result)
                logger.info(f"✓ AI categorized batch {batch_start // BATCH_SIZE + 1}: {len(batch_result)} items")

            except Exception as e:
                logger.warning(f"⚠️ AI categorization batch failed: {e}. Leaving as nao_categorizado.")

        # Apply categories back to transactions
        VALID_CATEGORIES = {
            "receita_vendas_produtos", "receita_servicos", "receita_locacao", "receita_comissoes",
            "receita_contratos_recorrentes", "impostos_sobre_vendas", "devolucoes", "descontos_concedidos",
            "cmv", "csp", "materia_prima", "insumos", "comissoes_sobre_vendas", "salarios_producao",
            "encargos_sociais_producao", "energia_producao", "manutencao_equipamentos_producao",
            "salarios_administrativos", "pro_labore", "encargos_sociais_administrativos", "aluguel",
            "condominio", "agua_energia", "material_escritorio", "honorarios_contabeis", "sistemas_softwares",
            "telefonia_internet", "marketing_publicidade", "propaganda_digital", "comissao_vendas", "fretes",
            "representantes_comerciais", "receita_financeira", "juros_ativos", "descontos_obtidos",
            "juros_passivos", "tarifas_bancarias", "iof", "multas_encargos", "irpj", "csll",
            "simples_nacional", "iptu", "taxas_municipais", "recuperacao_despesas", "venda_imobilizado",
            "indenizacoes_recebidas", "outras_receitas_eventuais", "perdas", "indenizacoes_pagas",
            "doacoes", "provisoes", "depreciacao", "amortizacao", "outras_despesas_operacionais",
            "nao_categorizado",
        }

        for txn in transactions:
            if (not txn.category or txn.category.strip().lower() in UNCATEGORIZED) and txn.description:
                ai_cat = description_to_category.get(txn.description, "nao_categorizado")
                txn.category = ai_cat if ai_cat in VALID_CATEGORIES else "nao_categorizado"

        return transactions

    def _extract_structured_data_from_excel(self, excel_text: str) -> FinancialDocument:
        """Extract structured data from Excel text using AI (with failover)"""
        return self._call_with_failover("excel", excel_text)

    def _extract_excel_with_openai(self, excel_text: str) -> FinancialDocument:
        """Extract structured data from Excel using OpenAI with caching and retry"""

        # Check cache
        cache_key = self._generate_cache_key(excel_text, "openai_excel")
        cached = self._get_cached_response(cache_key)
        if cached:
            return FinancialDocument(**cached)

        def _call_openai():
            logger.debug(
                f"🤖 Calling OpenAI {self.openai_model} for Excel extraction..."
            )

            # Check if model is a reasoning model (GPT-5, o1, o3 series)
            is_reasoning_model = any(model in self.openai_model.lower() for model in ['gpt-5', 'o1', 'o3'])

            api_params = {
                "model": self.openai_model,
                "messages": [
                    {
                        "role": "user",
                        "content": f"{self._get_extraction_prompt(self.user_company_info)}\n\nDADOS DO EXCEL:\n{excel_text}",
                    }
                ],
                "max_completion_tokens": 16000,  # Higher limit for reasoning models (reasoning + output)
                "store": False,  # CRITICAL: Do not store sensitive financial documents
            }

            # Add reasoning_effort for reasoning models to reduce unnecessary reasoning tokens
            if is_reasoning_model:
                api_params["reasoning_effort"] = "low"

            response = self._active_client.chat.completions.create(**api_params)
            return response

        try:
            response = self._call_with_retry(_call_openai)

            json_text = response.choices[0].message.content
            logger.debug("✓ OpenAI response received, parsing JSON...")
            data = json.loads(json_text)

            # Cache the response
            self._set_cached_response(cache_key, data)

            return FinancialDocument(**data)
        except Exception as e:
            logger.error(
                f"❌ OpenAI Excel extraction failed: {type(e).__name__}: {str(e)}"
            )
            raise

    def _extract_structured_data_from_pdf(self, pdf_base64: str) -> FinancialDocument:
        """Extract structured data from PDF using AI (with failover)"""
        return self._call_with_failover("pdf", pdf_base64)

    def _extract_structured_data(
        self, image_base64: str, image_type: str, page_num: int = None
    ) -> FinancialDocument:
        """Extract structured financial data using AI (with failover)"""
        return self._call_with_failover("image", image_base64, image_type)

    def _get_extraction_prompt(self, user_company_info: dict = None) -> str:
        """Get the prompt for structured data extraction with user context"""

        user_context = ""
        if user_company_info:
            user_context = f"""
USER COMPANY INFORMATION (This is YOUR company - use this to determine if transaction is income or expense):
- Company Name: {user_company_info.get('company_name', 'N/A')}
- Legal Name: {user_company_info.get('legal_name', 'N/A')}
- CNPJ: {user_company_info.get('cnpj', 'N/A')}

IMPORTANT: Compare the issuer and recipient with the user's company info above:
- If the USER is the ISSUER (document issued BY the user's company) → transaction_type: "income" (you sold/provided service)
- If the USER is the RECIPIENT (document issued TO the user's company) → transaction_type: "expense" (you bought/received service)
- Match by company name, legal name, or CNPJ
- transaction_type must ALWAYS be either "income" or "expense" - NEVER use "other" or any other value!
- If unclear, default to "expense"
"""

        # Build known items context for better categorization
        known_items_context = ""
        known_items = getattr(self, 'known_items', None)
        if known_items:
            # Sanitize user-editable fields before injecting into prompt
            def _sanitize(text: str, max_len: int = 100) -> str:
                if not text:
                    return ""
                # Strip control chars and prompt-injection markers
                clean = re.sub(r'[\x00-\x1f\x7f]', '', text)
                # Collapse whitespace
                clean = re.sub(r'\s+', ' ', clean).strip()
                return clean[:max_len]

            # Group by category for token efficiency
            by_category = {}
            for item in known_items:
                cat = _sanitize(item.get("category") or "sem_categoria", 50)
                if cat not in by_category:
                    by_category[cat] = []
                name = _sanitize(item.get("name", ""), 100)
                alias = _sanitize(item.get("alias", ""), 100)
                alias_str = f" → {alias}" if alias else ""
                count_str = f" (visto {item['times_appeared']}x)" if item.get("times_appeared", 0) > 1 else ""
                by_category[cat].append(f"  - {name}{alias_str}{count_str}")

            lines = []
            for cat, items_list in by_category.items():
                lines.append(f"\nCategoria: {cat}")
                lines.extend(items_list)

            known_items_context = f"""
KNOWN ITEMS FROM PREVIOUS UPLOADS:
The following items have been seen before in this organization's documents.
Use this information to correctly categorize matching items and assign the same category.
If an item matches a known item below, use the same category.
{"".join(lines)}
"""

        return f"""You are a financial document analyzer for Brazilian documents. Extract structured data from this document image.

This is a Brazilian financial document. It may be one of these types:
- Nota Fiscal (tax invoice)
- Nota Fiscal Eletrônica (NFe - electronic invoice)
- Nota Fiscal de Cancelamento (NFe cancellation document)
- Recibo (receipt)
- Cupom Fiscal (fiscal coupon/receipt from retail)
- Boleto (payment slip)
- Promissória (promissory note)
- Cheque (check)
- Relatório Financeiro (financial report/statement)
- Comprovante de Pagamento (payment proof)
- Outros documentos comerciais brasileiros

The document text is in Portuguese (pt-BR). Extract data accurately considering Brazilian formats.

{user_context}
{known_items_context}
Return ONLY a valid JSON object with the following structure (no markdown, no explanation):

{{
  "document_type": "invoice|receipt|boleto|check|promissory_note|fiscal_coupon|payment_proof|statement|expense|other",
  "document_number": "document number if present",
  "issue_date": "YYYY-MM-DD format",
  "transaction_type": "income|expense",
  "category": "one of the V2 category names listed below",
  "issuer": {{
    "name": "company/person who issued",
    "legal_name": "legal/official name if different",
    "tax_id": "CNPJ/CPF if present",
    "address": "full address if present",
    "phone": "phone if present",
    "email": "email if present"
  }},
  "recipient": {{
    "name": "company/person who received",
    "legal_name": "legal name if different",
    "tax_id": "CNPJ/CPF if present"
  }},
  "line_items": [
    {{
      "description": "item description",
      "quantity": 1.0,
      "unit_price": "10.50",
      "total_price": "10.50",
      "product_code": "code if present",
      "category": "one of the V2 category names listed below (REQUIRED - always assign a category per item)",
      "transaction_type": "income|expense"
    }}
  ],
  "subtotal": "100.00",
  "tax_amount": "10.00",
  "tax_rate": 10.0,
  "discount": "5.00",
  "total_amount": "105.00",
  "currency": "BRL",
  "payment_info": {{
    "status": "paid|unpaid|partial|pending",
    "method": "credit_card|cash|transfer|pix|boleto|other",
    "due_date": "YYYY-MM-DD",
    "payment_date": "YYYY-MM-DD if paid"
  }},
  "transactions": [
    {{
      "date": "YYYY-MM-DD",
      "description": "transaction description",
      "category": "one of the V2 category names (REQUIRED per transaction)",
      "amount": "100.00",
      "transaction_type": "income|expense",
      "counterparty": "supplier or client name if available"
    }}
  ],
  "is_cancellation": false,
  "original_document_number": null,
  "notes": "any additional notes or observations",
  "confidence_score": 0.95
}}

**CRITICAL - confidence_score field**:
- You MUST include a "confidence_score" field (0.0 to 1.0) indicating your confidence in the extraction
- Use 0.9-1.0 for clear, well-formatted documents with all data visible
- Use 0.7-0.9 for documents with some unclear text but mostly readable
- Use 0.5-0.7 for documents with significant quality issues or missing data
- Use 0.0-0.5 for documents you cannot read properly or are highly uncertain about

Important rules for Brazilian documents:
- Document is in Portuguese (pt-BR)
- Use null for missing fields (except payment_info.status - use "pending" if unclear)
- **CRITICAL - Brazilian Number Format** (READ THIS CAREFULLY):
  * In Brazil, the number format is OPPOSITE of US format!
  * COMMA (,) = DECIMAL SEPARATOR
  * PERIOD (.) = THOUSANDS SEPARATOR
  * Examples of what you'll see in the document:
    - "38.000,00" = thirty-eight thousand reais (NOT thirty-eight reais!) → convert to "38000.00"
    - "1.234,56" = one thousand two hundred thirty-four and 56 cents → convert to "1234.56"
    - "4.500,00" = four thousand five hundred → convert to "4500.00"
    - "1.600.000,00" = one million six hundred thousand → convert to "1600000.00"
    - "38,00" = thirty-eight reais → convert to "38.00"
  * COMMON MISTAKE TO AVOID: If you see "38.000,00" DO NOT interpret as "38.00"! It's "38000.00"!
  * You MUST convert to standard format for JSON (period for decimal, no thousands separator)
- Monetary values: strings with 2 decimal places in standard format (e.g., "1250.50")
- Dates: convert to YYYY-MM-DD format (from dd/mm/yyyy if needed)
- CNPJ format: XX.XXX.XXX/XXXX-XX (14 digits)
- CPF format: XXX.XXX.XXX-XX (11 digits)
- **CATEGORY FIELD** - Use EXACTLY one of these V2 category names (Plano de Contas):
  RECEITA (Revenue): "receita_vendas_produtos", "receita_servicos", "receita_locacao", "receita_comissoes", "receita_contratos_recorrentes"
  DEDUÇÕES (Deductions): "impostos_sobre_vendas", "devolucoes", "descontos_concedidos"
  CUSTOS VARIÁVEIS (Variable Costs): "cmv", "csp", "materia_prima", "insumos", "comissoes_sobre_vendas"
  CUSTOS FIXOS PRODUÇÃO (Fixed Production Costs): "salarios_producao", "encargos_sociais_producao", "energia_producao", "manutencao_equipamentos_producao"
  DESPESAS ADMIN (Admin Expenses): "salarios_administrativos", "pro_labore", "encargos_sociais_administrativos", "aluguel", "condominio", "agua_energia", "material_escritorio", "honorarios_contabeis", "sistemas_softwares", "telefonia_internet"
  DESPESAS COMERCIAIS (Commercial): "marketing_publicidade", "propaganda_digital", "comissao_vendas", "fretes", "representantes_comerciais"
  FINANCEIRO (Financial): "receita_financeira", "juros_ativos", "descontos_obtidos", "juros_passivos", "tarifas_bancarias", "iof", "multas_encargos"
  TRIBUTOS (Taxes): "irpj", "csll", "simples_nacional", "iptu", "taxas_municipais"
  OUTRAS (Other): "recuperacao_despesas", "venda_imobilizado", "indenizacoes_recebidas", "outras_receitas_eventuais", "perdas", "indenizacoes_pagas", "doacoes", "provisoes", "depreciacao", "amortizacao", "outras_despesas_operacionais"
  FALLBACK: "nao_categorizado" (use ONLY as absolute last resort - try hard to match a real category first)
- **CRITICAL - CATEGORY IS MANDATORY**: You MUST assign a category to EVERY item and to the document itself. Never leave category as null or empty. Analyze each item's description and assign the most appropriate category from the list above. For bank statements, PIX transfers, card purchases etc., categorize based on the description (e.g., card purchase → "outras_despesas_operacionais", bank fees → "tarifas_bancarias", PIX sent → analyze context). Each line_item MUST have its own "category" field.
- Common payment methods: "pix", "boleto", "cartao_credito", "transferencia", "dinheiro"
- **EXCEL/SPREADSHEET WITH MULTIPLE TRANSACTIONS**: If the document is an Excel spreadsheet, bank statement, or ledger with multiple transactions (one per row), use the "transactions" array (NOT "line_items"). Each transaction should have its own date, description, category, amount, transaction_type, and counterparty (supplier/client name if present in the data). If each row has a different supplier/client, capture that in the counterparty field of each transaction.
- **INVOICES WITH PRODUCTS**: For invoices/receipts with product line items, use "line_items" array.
- **CRITICAL - LINE ITEMS**: You MUST extract EVERY SINGLE line item from the document
  * Do NOT skip items or summarize - include ALL items even if there are many
  * If you see a table with products/services, extract EVERY row
  * Common invoices have 5-50+ items - extract them ALL
  * Missing items = wrong total amount = failed extraction
- Preserve Portuguese text in descriptions
- Be thorough and accurate with Brazilian tax information
- **NFe CANCELLATION DETECTION**: If the document is a cancellation notice (cancelamento de NFe, carta de correção cancelando NF, evento de cancelamento):
  * Set "is_cancellation": true
  * Set "original_document_number" to the number of the original NF being cancelled
  * Look for keywords: "CANCELAMENTO", "CANCELADA", "CARTA DE CORREÇÃO", "INUTILIZAÇÃO", "Evento: Cancelamento"
  * The original NF number is often referenced as "Chave de Acesso" or "NFe ref"
  * Set total_amount to "0.00" since cancelled NFs have no financial effect
- Return ONLY the JSON, no other text"""

    def _extract_pdf_with_openai(self, pdf_base64: str) -> FinancialDocument:
        """
        Extract structured data from PDF using OpenAI (native PDF support)

        Uses data:application/pdf;base64,{pdf_base64} format
        OpenAI extracts both text and images from PDF internally
        """

        # Check cache first
        cache_key = self._generate_cache_key(pdf_base64, "openai_pdf")
        cached = self._get_cached_response(cache_key)
        if cached:
            return FinancialDocument(**cached)

        def _call_openai():
            logger.debug(f"🤖 Calling OpenAI {self.openai_model} for PDF extraction...")

            # Check if model is a reasoning model (GPT-5, o1, o3 series)
            is_reasoning_model = any(model in self.openai_model.lower() for model in ['gpt-5', 'o1', 'o3'])

            api_params = {
                "model": self.openai_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": self._get_extraction_prompt(self.user_company_info)},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:application/pdf;base64,{pdf_base64}",
                                    "detail": "high"  # High detail for better text extraction
                                },
                            },
                        ],
                    }
                ],
                "max_completion_tokens": 16000,  # Higher limit for reasoning models
                "store": False,  # CRITICAL: Do not store sensitive financial documents
            }

            # Add reasoning_effort for reasoning models
            if is_reasoning_model:
                api_params["reasoning_effort"] = "low"

            response = self._active_client.chat.completions.create(**api_params)
            return response

        try:
            # Call with retry logic
            response = self._call_with_retry(_call_openai)

            json_text = response.choices[0].message.content
            logger.debug(f"✓ OpenAI response received, content length: {len(json_text) if json_text else 0}")

            if not json_text:
                logger.error(f"❌ OpenAI returned empty response. Full response: {response}")
                raise ValueError("OpenAI returned empty content")

            logger.debug("Parsing JSON...")
            data = json.loads(json_text)

            # Cache the response
            self._set_cached_response(cache_key, data)

            return FinancialDocument(**data)
        except Exception as e:
            logger.error(f"❌ OpenAI PDF extraction failed: {type(e).__name__}: {str(e)}")
            raise

    def _extract_with_openai(
        self, image_base64: str, image_type: str
    ) -> FinancialDocument:
        """Extract structured data using OpenAI with caching and retry logic"""

        # Check cache first
        cache_key = self._generate_cache_key(image_base64, "openai_image")
        cached = self._get_cached_response(cache_key)
        if cached:
            return FinancialDocument(**cached)

        def _call_openai():
            logger.debug(f"🤖 Calling OpenAI {self.openai_model} for extraction...")

            # Check if model is a reasoning model (GPT-5, o1, o3 series)
            is_reasoning_model = any(model in self.openai_model.lower() for model in ['gpt-5', 'o1', 'o3'])

            api_params = {
                "model": self.openai_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": self._get_extraction_prompt(self.user_company_info)},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/{image_type};base64,{image_base64}"
                                },
                            },
                        ],
                    }
                ],
                "max_completion_tokens": 16000,  # Higher limit for reasoning models (reasoning + output)
                "store": False,  # CRITICAL: Do not store sensitive financial documents
            }

            # Add reasoning_effort for reasoning models to reduce unnecessary reasoning tokens
            if is_reasoning_model:
                api_params["reasoning_effort"] = "low"

            response = self._active_client.chat.completions.create(**api_params)
            return response

        try:
            # Call with retry logic
            response = self._call_with_retry(_call_openai)

            json_text = response.choices[0].message.content
            logger.debug(f"✓ OpenAI response received, content length: {len(json_text) if json_text else 0}")

            if not json_text:
                logger.error(f"❌ OpenAI returned empty response. Full response: {response}")
                raise ValueError("OpenAI returned empty content")

            logger.debug("Parsing JSON...")
            data = json.loads(json_text)

            # Cache the response
            self._set_cached_response(cache_key, data)

            return FinancialDocument(**data)
        except Exception as e:
            logger.error(f"❌ OpenAI extraction failed: {type(e).__name__}: {str(e)}")
            raise

    # ------------------------------------------------------------------
    # Gemini extraction methods
    # ------------------------------------------------------------------

    def _extract_with_gemini(self, image_base64: str, image_type: str) -> FinancialDocument:
        """Extract financial data from image using Google Gemini"""
        cache_key = self._generate_cache_key(image_base64, "gemini_image")
        cached = self._get_cached_response(cache_key)
        if cached:
            return FinancialDocument(**cached)

        def _call_gemini():
            image_bytes = base64.b64decode(image_base64)
            image = Image.open(io.BytesIO(image_bytes))

            prompt = self._get_extraction_prompt(self.user_company_info)

            model = self._active_client  # GenerativeModel instance
            response = model.generate_content(
                [prompt, image],
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=16000,
                    temperature=0.1,
                ),
            )
            return response

        response = self._call_with_retry(_call_gemini)
        json_text = response.text

        # Clean markdown fences if present
        if json_text.startswith("```"):
            json_text = json_text.split("```")[1]
            if json_text.startswith("json"):
                json_text = json_text[4:]
        json_text = json_text.strip()

        data = json.loads(json_text)
        self._set_cached_response(cache_key, data)
        return FinancialDocument(**data)

    def _extract_pdf_with_gemini(self, pdf_base64: str) -> FinancialDocument:
        """Extract financial data from PDF using Google Gemini (native PDF support)"""
        cache_key = self._generate_cache_key(pdf_base64[:200], "gemini_pdf")
        cached = self._get_cached_response(cache_key)
        if cached:
            return FinancialDocument(**cached)

        def _call_gemini():
            pdf_bytes = base64.b64decode(pdf_base64)

            prompt = self._get_extraction_prompt(self.user_company_info)

            model = self._active_client
            response = model.generate_content(
                [
                    prompt,
                    {"mime_type": "application/pdf", "data": pdf_bytes},
                ],
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=16000,
                    temperature=0.1,
                ),
            )
            return response

        response = self._call_with_retry(_call_gemini)
        json_text = response.text

        if json_text.startswith("```"):
            json_text = json_text.split("```")[1]
            if json_text.startswith("json"):
                json_text = json_text[4:]
        json_text = json_text.strip()

        data = json.loads(json_text)
        self._set_cached_response(cache_key, data)
        return FinancialDocument(**data)

    def _extract_excel_with_gemini(self, excel_text: str) -> FinancialDocument:
        """Extract financial data from Excel text using Google Gemini"""
        cache_key = self._generate_cache_key(excel_text[:500], "gemini_excel")
        cached = self._get_cached_response(cache_key)
        if cached:
            return FinancialDocument(**cached)

        def _call_gemini():
            prompt = f"{self._get_extraction_prompt(self.user_company_info)}\n\nDADOS DO EXCEL:\n{excel_text}"

            model = self._active_client
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=16000,
                    temperature=0.1,
                ),
            )
            return response

        response = self._call_with_retry(_call_gemini)
        json_text = response.text

        if json_text.startswith("```"):
            json_text = json_text.split("```")[1]
            if json_text.startswith("json"):
                json_text = json_text[4:]
        json_text = json_text.strip()

        data = json.loads(json_text)
        self._set_cached_response(cache_key, data)
        return FinancialDocument(**data)

    def _ai_detect_columns_gemini(self, df) -> dict:
        """Use Gemini to detect column structure in tabular data"""
        prompt = self._build_columns_prompt(df)

        def _call_gemini():
            model = self._active_client
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=4000,
                    temperature=0.1,
                ),
            )
            return response

        response = self._call_with_retry(_call_gemini)
        json_text = response.text

        if json_text.startswith("```"):
            json_text = json_text.split("```")[1]
            if json_text.startswith("json"):
                json_text = json_text[4:]

        return json.loads(json_text.strip())

    # ------------------------------------------------------------------
    # Nova / Bedrock extraction methods
    # ------------------------------------------------------------------

    def _extract_with_nova(self, image_base64: str, image_type: str) -> FinancialDocument:
        """Extract financial data from image using Amazon Nova via Bedrock"""
        cache_key = self._generate_cache_key(image_base64, "nova_image")
        cached = self._get_cached_response(cache_key)
        if cached:
            return FinancialDocument(**cached)

        # Map common types to Bedrock format
        format_map = {"jpeg": "jpeg", "jpg": "jpeg", "png": "png", "webp": "webp", "gif": "gif"}
        img_format = format_map.get(image_type.lower(), "png")

        def _call_nova():
            prompt = self._get_extraction_prompt(self.user_company_info)

            bedrock = self._active_client  # boto3 bedrock-runtime client
            response = bedrock.converse(
                modelId=self.nova_model,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "image": {
                                "format": img_format,
                                "source": {"bytes": base64.b64decode(image_base64)},
                            }
                        },
                        {"text": prompt},
                    ],
                }],
                inferenceConfig={"maxTokens": 16000, "temperature": 0.1},
            )
            return response

        response = self._call_with_retry(_call_nova)
        json_text = response["output"]["message"]["content"][0]["text"]

        if json_text.startswith("```"):
            json_text = json_text.split("```")[1]
            if json_text.startswith("json"):
                json_text = json_text[4:]
        json_text = json_text.strip()

        data = json.loads(json_text)
        self._set_cached_response(cache_key, data)
        return FinancialDocument(**data)

    def _extract_pdf_with_nova(self, pdf_base64: str) -> FinancialDocument:
        """Extract financial data from PDF using Amazon Nova.
        Nova doesn't natively support PDF in Converse API - convert to image first."""
        # It's base64 - write to temp file first
        import tempfile

        pdf_bytes = base64.b64decode(pdf_base64)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.write(pdf_bytes)
        tmp.close()

        try:
            poppler_path = os.getenv("POPPLER_PATH") or None
            images = convert_from_path(tmp.name, first_page=1, last_page=1, dpi=200,
                                       poppler_path=poppler_path)
            if not images:
                raise ValueError("Failed to convert PDF to image for Nova processing")

            img_buffer = io.BytesIO()
            images[0].save(img_buffer, format="PNG")
            img_base64 = base64.b64encode(img_buffer.getvalue()).decode("utf-8")

            return self._extract_with_nova(img_base64, "png")
        finally:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

    def _extract_excel_with_nova(self, excel_text: str) -> FinancialDocument:
        """Extract financial data from Excel text using Amazon Nova via Bedrock"""
        cache_key = self._generate_cache_key(excel_text[:500], "nova_excel")
        cached = self._get_cached_response(cache_key)
        if cached:
            return FinancialDocument(**cached)

        def _call_nova():
            prompt = f"{self._get_extraction_prompt(self.user_company_info)}\n\nDADOS DO EXCEL:\n{excel_text}"

            bedrock = self._active_client
            response = bedrock.converse(
                modelId=self.nova_model,
                messages=[{
                    "role": "user",
                    "content": [{"text": prompt}],
                }],
                inferenceConfig={"maxTokens": 16000, "temperature": 0.1},
            )
            return response

        response = self._call_with_retry(_call_nova)
        json_text = response["output"]["message"]["content"][0]["text"]

        if json_text.startswith("```"):
            json_text = json_text.split("```")[1]
            if json_text.startswith("json"):
                json_text = json_text[4:]
        json_text = json_text.strip()

        data = json.loads(json_text)
        self._set_cached_response(cache_key, data)
        return FinancialDocument(**data)

    def _ai_detect_columns_nova(self, df) -> dict:
        """Use Amazon Nova to detect column structure in tabular data"""
        prompt = self._build_columns_prompt(df)

        def _call_nova():
            bedrock = self._active_client
            response = bedrock.converse(
                modelId=self.nova_model,
                messages=[{
                    "role": "user",
                    "content": [{"text": prompt}],
                }],
                inferenceConfig={"maxTokens": 4000, "temperature": 0.1},
            )
            return response

        response = self._call_with_retry(_call_nova)
        json_text = response["output"]["message"]["content"][0]["text"]

        if json_text.startswith("```"):
            json_text = json_text.split("```")[1]
            if json_text.startswith("json"):
                json_text = json_text[4:]

        return json.loads(json_text.strip())

    def _process_xml(self, xml_path: Path) -> dict:
        """
        Process XML files (Brazilian fiscal documents: NFe, NFSe, CTe, etc.)

        Supports:
        - NFe (Nota Fiscal Eletrônica) - Brazilian invoice
        - NFSe (Nota Fiscal de Serviço) - Service invoice
        - CTe (Conhecimento de Transporte) - Transport document
        - Generic XML with AI fallback
        """
        logger.info(f"📄 Processing XML: {xml_path.name}")

        if not XML_AVAILABLE:
            return {
                "file_name": xml_path.name,
                "file_type": "xml",
                "status": "failed",
                "error": "XML processing libraries not available. Install lxml and xmltodict.",
            }

        try:
            # Read and parse XML
            with open(xml_path, "r", encoding="utf-8") as f:
                xml_content = f.read()

            xml_dict = xmltodict.parse(xml_content)

            # Detect XML type and extract data
            extracted_data = None

            # Try NFe (Nota Fiscal Eletrônica)
            if "nfeProc" in xml_dict or "NFe" in xml_dict:
                logger.info("Detected NFe (Nota Fiscal Eletrônica)")
                extracted_data = self._extract_nfe(xml_dict)

            # Try NFSe (Nota Fiscal de Serviço)
            elif "CompNfse" in xml_dict or "NFSe" in xml_dict:
                logger.info("Detected NFSe (Nota Fiscal de Serviço)")
                extracted_data = self._extract_nfse(xml_dict)

            # Try CTe (Conhecimento de Transporte)
            elif "cteProc" in xml_dict or "CTe" in xml_dict:
                logger.info("Detected CTe (Conhecimento de Transporte)")
                extracted_data = self._extract_cte(xml_dict)

            # Item 4: Handle cancellation documents (extract cancellation data instead of rejecting)
            elif "procCancNFe" in xml_dict or "retCancNFe" in xml_dict:
                logger.info("Detected NFe cancellation document - extracting cancellation data")
                extracted_data = self._extract_nfe_cancellation(xml_dict)

            # Handle cancellation events (procEventoNFe with tpEvento=110111)
            elif "procEventoNFe" in xml_dict:
                evento = xml_dict.get("procEventoNFe", {}).get("evento", {})
                inf_evento = evento.get("infEvento", {})
                tp_evento = inf_evento.get("tpEvento", "")
                if tp_evento == "110111":  # Cancellation event
                    logger.info("Detected NFe cancellation event - extracting cancellation data")
                    extracted_data = self._extract_nfe_cancellation_event(xml_dict)
                else:
                    logger.warning(f"Unknown NFe event type: {tp_evento}")
                    return {
                        "file_name": xml_path.name,
                        "file_type": "xml",
                        "status": "failed",
                        "error": f"Tipo de evento NFe não suportado: {tp_evento}",
                    }

            # Generic XML - not supported
            else:
                logger.warning(f"Unknown XML format in {xml_path.name}")
                return {
                    "file_name": xml_path.name,
                    "file_type": "xml",
                    "status": "failed",
                    "error": "Formato XML não reconhecido. Formatos suportados: NFe, NFSe, CTe",
                }

            return {
                "file_name": xml_path.name,
                "file_type": "xml",
                "status": "success",
                "extracted_data": extracted_data,
            }

        except Exception as e:
            logger.error(f"❌ XML processing failed: {type(e).__name__}: {str(e)}")
            return {
                "file_name": xml_path.name,
                "file_type": "xml",
                "status": "failed",
                "error": str(e),
            }

    def _extract_nfe(self, xml_dict: dict) -> FinancialDocument:
        """Extract data from NFe (Brazilian electronic invoice)"""
        try:
            # Navigate NFe structure
            nfe_root = xml_dict.get("nfeProc", xml_dict.get("NFe", {}))
            nfe = nfe_root.get("NFe", nfe_root)
            inf_nfe = nfe.get("infNFe", {})

            # Get issuer info (emitente)
            emit = inf_nfe.get("emit", {})

            # Get recipient info (destinatário)
            dest = inf_nfe.get("dest", {})

            # Get items
            det = inf_nfe.get("det", [])
            if not isinstance(det, list):
                det = [det]

            # Get totals
            total = inf_nfe.get("total", {}).get("ICMSTot", {})

            # Determine transaction type by comparing issuer with user's company
            issuer_cnpj = (emit.get("CNPJ") or emit.get("CPF") or "").strip()
            recipient_cnpj = (dest.get("CNPJ") or dest.get("CPF") or "").strip()

            user_cnpj = ""
            user_name = ""
            if hasattr(self, 'user_company_info') and self.user_company_info:
                user_cnpj = (self.user_company_info.get("cnpj") or "").replace(".", "").replace("/", "").replace("-", "").strip()
                user_name = (self.user_company_info.get("company_name") or "").lower().strip()

            # Clean CNPJs for comparison (remove formatting)
            clean_issuer = issuer_cnpj.replace(".", "").replace("/", "").replace("-", "")
            clean_recipient = recipient_cnpj.replace(".", "").replace("/", "").replace("-", "")

            # Decision: if user is the ISSUER → receita (they sold), if RECIPIENT → despesa (they bought)
            if user_cnpj and clean_issuer == user_cnpj:
                txn_type = "receita"
                txn_category = "receita_vendas_produtos"
                logger.info(f"📋 NFe: user is ISSUER (CNPJ match) → receita")
            elif user_cnpj and clean_recipient == user_cnpj:
                txn_type = "despesa"
                txn_category = "compras_mercadorias"
                logger.info(f"📋 NFe: user is RECIPIENT (CNPJ match) → despesa")
            elif user_name and user_name in (emit.get("xNome", "").lower()):
                txn_type = "receita"
                txn_category = "receita_vendas_produtos"
                logger.info(f"📋 NFe: user is ISSUER (name match) → receita")
            elif user_name and user_name in (dest.get("xNome", "").lower()):
                txn_type = "despesa"
                txn_category = "compras_mercadorias"
                logger.info(f"📋 NFe: user is RECIPIENT (name match) → despesa")
            else:
                # Default: assume user is issuer (they uploaded their own NFe)
                txn_type = "receita"
                txn_category = "receita_vendas_produtos"
                logger.info(f"📋 NFe: no CNPJ match, defaulting to receita (issuer)")

            # Build FinancialDocument
            return FinancialDocument(
                document_type="invoice",
                document_number=inf_nfe.get("@Id", "").replace("NFe", ""),
                issue_date=self._parse_date(inf_nfe.get("ide", {}).get("dhEmi", "")),
                transaction_type=txn_type,
                category=txn_category,
                issuer={
                    "name": emit.get("xNome", ""),
                    "tax_id": emit.get("CNPJ", emit.get("CPF", "")),
                    "address": f"{emit.get('enderEmit', {}).get('xLgr', '')}, {emit.get('enderEmit', {}).get('nro', '')}, {emit.get('enderEmit', {}).get('xBairro', '')}, {emit.get('enderEmit', {}).get('xMun', '')} - {emit.get('enderEmit', {}).get('UF', '')}",
                },
                recipient={
                    "name": dest.get("xNome", ""),
                    "tax_id": dest.get("CNPJ", dest.get("CPF", "")),
                },
                line_items=[
                    {
                        "description": item.get("prod", {}).get("xProd", ""),
                        "quantity": float(self._parse_brazilian_number(item.get("prod", {}).get("qCom", "0"))),
                        "unit_price": self._parse_brazilian_number(item.get("prod", {}).get("vUnCom", "0")),
                        "total_price": self._parse_brazilian_number(item.get("prod", {}).get("vProd", "0")),
                    }
                    for item in det
                ],
                subtotal=self._parse_brazilian_number(total.get("vProd", "0")),
                tax_amount=self._parse_brazilian_number(total.get("vTotTrib", "0")),
                total_amount=self._parse_brazilian_number(total.get("vNF", "0")),
                currency="BRL",
                confidence_score=0.98,  # High confidence for structured XML data
            )
        except Exception as e:
            logger.error(f"❌ NFe extraction failed: {e}")
            import traceback
            logger.error(f"   Traceback:\n{traceback.format_exc()}")
            # Return empty document instead of trying to send XML as image
            return FinancialDocument(
                document_type="invoice",
                status="failed",
                error=f"Failed to extract NFe data: {str(e)}",
                currency="BRL"
            )

    def _extract_nfse(self, xml_dict: dict) -> FinancialDocument:
        """Extract data from NFSe (Brazilian service invoice)"""
        # NFSe has many different layouts depending on the municipality
        # This is a simplified version - may need customization
        try:
            comp_nfse = xml_dict.get("CompNfse", {})
            nfse = comp_nfse.get("Nfse", comp_nfse.get("NFSe", {}))
            inf_nfse = nfse.get("InfNfse", nfse)

            # Determine transaction type by comparing with user company
            prestador = inf_nfse.get("PrestadorServico", {}) or inf_nfse.get("Prestador", {})
            tomador = inf_nfse.get("TomadorServico", {}) or inf_nfse.get("Tomador", {})
            prestador_cnpj = (prestador.get("IdentificacaoPrestador", {}).get("Cnpj", "") or
                              prestador.get("IdentificacaoPrestador", {}).get("CpfCnpj", {}).get("Cnpj", "")).strip()
            tomador_cnpj = (tomador.get("IdentificacaoTomador", {}).get("CpfCnpj", {}).get("Cnpj", "") or "").strip()

            user_cnpj = ""
            if hasattr(self, 'user_company_info') and self.user_company_info:
                user_cnpj = (self.user_company_info.get("cnpj") or "").replace(".", "").replace("/", "").replace("-", "").strip()

            clean_prest = prestador_cnpj.replace(".", "").replace("/", "").replace("-", "")
            clean_tom = tomador_cnpj.replace(".", "").replace("/", "").replace("-", "")

            if user_cnpj and clean_prest == user_cnpj:
                txn_type = "receita"
                txn_category = "receita_servicos"
            elif user_cnpj and clean_tom == user_cnpj:
                txn_type = "despesa"
                txn_category = "servicos_contratados"
            else:
                txn_type = "receita"  # Default: assume user is prestador
                txn_category = "receita_servicos"

            return FinancialDocument(
                document_type="invoice",
                document_number=inf_nfse.get("Numero", ""),
                issue_date=self._parse_date(inf_nfse.get("DataEmissao", "")),
                transaction_type=txn_type,
                category=txn_category,
                total_amount=str(
                    inf_nfse.get("Servico", {})
                    .get("Valores", {})
                    .get("ValorServicos", "0")
                ),
                currency="BRL",
                confidence_score=0.98,  # High confidence for structured XML data
            )
        except Exception as e:
            logger.warning(f"NFSe extraction failed, falling back to AI: {e}")
            xml_text = json.dumps(xml_dict, indent=2, ensure_ascii=False)
            result = self._extract_structured_data(xml_text, "xml")
            return result.get("extracted_data", FinancialDocument())

    def _extract_cte(self, xml_dict: dict) -> FinancialDocument:
        """Extract data from CTe (Brazilian transport document)"""
        try:
            cte_root = xml_dict.get("cteProc", xml_dict.get("CTe", {}))
            cte = cte_root.get("CTe", cte_root)
            inf_cte = cte.get("infCte", {})

            return FinancialDocument(
                document_type="invoice",
                document_number=inf_cte.get("@Id", "").replace("CTe", ""),
                issue_date=self._parse_date(inf_cte.get("ide", {}).get("dhEmi", "")),
                transaction_type="expense",  # CTe is typically transport expense
                category="fretes",
                total_amount=str(inf_cte.get("vPrest", {}).get("vTPrest", "0")),
                currency="BRL",
                confidence_score=0.98,  # High confidence for structured XML data
            )
        except Exception as e:
            logger.warning(f"CTe extraction failed, falling back to AI: {e}")
            xml_text = json.dumps(xml_dict, indent=2, ensure_ascii=False)
            result = self._extract_structured_data(xml_text, "xml")
            return result.get("extracted_data", FinancialDocument())

    def _extract_nfe_cancellation(self, xml_dict: dict) -> FinancialDocument:
        """
        Extract data from NFe cancellation document (procCancNFe/retCancNFe).
        Item 4: Instead of rejecting, we extract the original NF number and mark as cancellation.
        """
        try:
            canc_root = xml_dict.get("procCancNFe", xml_dict.get("retCancNFe", {}))

            # Try to get the cancelled NF key (chave de acesso)
            inf_canc = canc_root.get("cancNFe", canc_root).get("infCanc", canc_root.get("infCanc", {}))
            chave = inf_canc.get("chNFe", "")
            protocolo = inf_canc.get("nProt", "")
            data_canc = inf_canc.get("dhRecbto", inf_canc.get("dhRegEvento", ""))

            # Use the full chave de acesso as original_document_number
            # since that's what _extract_nfe stores as document_number
            return FinancialDocument(
                document_type="invoice",
                document_number=protocolo or chave,
                issue_date=self._parse_date(data_canc),
                transaction_type="expense",
                category="nao_categorizado",
                total_amount="0.00",
                currency="BRL",
                is_cancellation=True,
                original_document_number=chave,
                notes=f"Cancelamento de NFe. Chave: {chave}",
                confidence_score=0.99,
            )
        except Exception as e:
            logger.error(f"NFe cancellation extraction failed: {e}")
            return FinancialDocument(
                document_type="invoice",
                total_amount="0.00",
                currency="BRL",
                is_cancellation=True,
                notes=f"Cancellation document (extraction error: {str(e)})",
                confidence_score=0.5,
            )

    def _extract_nfe_cancellation_event(self, xml_dict: dict) -> FinancialDocument:
        """
        Extract data from NFe cancellation event (procEventoNFe with tpEvento=110111).
        Item 4: Handles the newer event-based cancellation format.
        """
        try:
            proc_evento = xml_dict.get("procEventoNFe", {})
            evento = proc_evento.get("evento", {})
            inf_evento = evento.get("infEvento", {})

            chave = inf_evento.get("chNFe", "")
            data_evento = inf_evento.get("dhEvento", "")
            det_evento = inf_evento.get("detEvento", {})
            justificativa = det_evento.get("xJust", "Cancelamento")
            protocolo = det_evento.get("nProt", inf_evento.get("nSeqEvento", ""))

            # Use the full chave de acesso as original_document_number
            # since that's what _extract_nfe stores as document_number
            return FinancialDocument(
                document_type="invoice",
                document_number=protocolo or chave,
                issue_date=self._parse_date(data_evento),
                transaction_type="expense",
                category="nao_categorizado",
                total_amount="0.00",
                currency="BRL",
                is_cancellation=True,
                original_document_number=chave,
                notes=f"Cancelamento de NFe. Justificativa: {justificativa}. Chave: {chave}",
                confidence_score=0.99,
            )
        except Exception as e:
            logger.error(f"NFe cancellation event extraction failed: {e}")
            return FinancialDocument(
                document_type="invoice",
                total_amount="0.00",
                currency="BRL",
                is_cancellation=True,
                notes=f"Cancellation event (extraction error: {str(e)})",
                confidence_score=0.5,
            )

    def _parse_date(self, date_str: str) -> str:
        """Parse Brazilian XML date formats to YYYY-MM-DD"""
        if not date_str:
            return ""

        try:
            # Handle datetime format (2024-01-15T10:30:00-03:00)
            if "T" in date_str:
                date_str = date_str.split("T")[0]

            # Already in YYYY-MM-DD format
            if len(date_str) == 10 and "-" in date_str:
                return date_str

            # Handle YYYYMMDD format
            if len(date_str) == 8 and date_str.isdigit():
                return f"{date_str[0:4]}-{date_str[4:6]}-{date_str[6:8]}"

            return date_str
        except:
            return date_str

    def _parse_cell_as_decimal(self, val) -> Decimal:
        """
        Robustly parse a cell value into a Decimal.
        Handles:
          - Already numeric (float/int from pandas)
          - Brazilian format strings: "1.234,56" or "-1.234,56"
          - Currency prefixed: "R$ 1.234,56"
          - Parentheses for negatives: "(1.234,56)"
          - Standard format: "1234.56"
        Returns Decimal or raises ValueError if unparseable.
        """
        import pandas as pd

        if pd.isna(val):
            raise ValueError("NaN value")

        # If already numeric (pandas parsed it correctly), just convert
        if isinstance(val, (int, float)):
            return Decimal(str(val))

        # String processing
        val_str = str(val).strip()
        if not val_str:
            raise ValueError("Empty string")

        # Remove currency symbols and whitespace
        import re
        val_str = re.sub(r'[R$\s]', '', val_str)

        # Handle parentheses as negative: (1.234,56) -> -1.234,56
        if val_str.startswith('(') and val_str.endswith(')'):
            val_str = '-' + val_str[1:-1]

        # Use existing Brazilian number parser (returns standard format string)
        standard = self._parse_brazilian_number(val_str)

        # Final conversion
        result = Decimal(standard)
        return result

    def _parse_brazilian_number(self, value: str) -> str:
        """
        Parse Brazilian number format to standard format
        Brazilian format: 1.234,56 (period = thousands, comma = decimal)
        Standard format: 1234.56
        """
        if not value:
            return "0"

        try:
            # Convert to string if needed
            value_str = str(value)

            # Remove any spaces
            value_str = value_str.strip().replace(" ", "")

            # If already in standard format (only digits and one period as decimal)
            # Example: "1234.56"
            if "," not in value_str and value_str.replace(".", "").replace("-", "").isdigit():
                # Check if period is decimal separator (2 digits after)
                if "." in value_str:
                    parts = value_str.split(".")
                    if len(parts) == 2 and len(parts[1]) == 2:
                        return value_str  # Already standard format
                return value_str

            # Brazilian format: has comma as decimal separator
            # Example: "1.234,56" or "4.500,00"
            if "," in value_str:
                # Remove thousands separator (periods)
                value_str = value_str.replace(".", "")
                # Replace decimal separator (comma) with period
                value_str = value_str.replace(",", ".")
                return value_str

            # If no comma and multiple periods, it's thousands separator only
            # Example: "1.234" should be "1234"
            if value_str.count(".") > 1:
                value_str = value_str.replace(".", "")
                return value_str

            return value_str

        except Exception as e:
            logger.warning(f"Failed to parse number '{value}': {e}")
            return str(value)

    def _ai_detect_columns(self, df) -> dict:
        """Use AI to detect column purposes when keyword matching fails.
        Routes through _call_with_failover for multi-key + provider failover.
        """
        try:
            return self._call_with_failover("columns", df)
        except Exception as e:
            logger.warning(f"AI column detection failed: {e}")
            return {}

    def _build_columns_prompt(self, df) -> str:
        """Build the prompt for AI column detection"""
        column_info = []
        for col in df.columns[:15]:  # Limit to first 15 columns to save tokens
            samples = df[col].dropna().head(3).tolist()
            sample_str = ", ".join(str(s)[:50] for s in samples)
            column_info.append(f"Column '{col}': {sample_str}")

        columns_text = "\n".join(column_info)

        return f"""Analyze this spreadsheet and identify which columns contain:
- date: Transaction or document date
- amount: Monetary value/amount (may be called: valor, competencia, money, value, débito, crédito, etc.)
- description: Transaction/item description
- category: Transaction category/type

Columns and sample data:
{columns_text}

Respond ONLY with a JSON object mapping purposes to exact column names:
{{"date": "COLUMN_NAME", "amount": "COLUMN_NAME", "description": "COLUMN_NAME", "category": "COLUMN_NAME"}}

Use null for any column type you cannot identify."""

    def _ai_detect_columns_openai(self, df) -> dict:
        """AI column detection using OpenAI API format"""
        prompt = self._build_columns_prompt(df)

        def _call():
            response = self._active_client.chat.completions.create(
                model=self.openai_model,
                messages=[{"role": "user", "content": prompt}],
                max_completion_tokens=200,
                store=False,
            )
            return response.choices[0].message.content.strip()

        result_text = self._call_with_retry(_call)
        if "```json" in result_text:
            result_text = result_text.split("```json")[1].split("```")[0].strip()
        elif "```" in result_text:
            result_text = result_text.split("```")[1].split("```")[0].strip()
        return json.loads(result_text)

    # ========================================================================
    # V2 File Format Processors (Item 2 - partner feedback)
    # OFX, OFC, DOCX, TXT
    # ========================================================================

    def _process_ofx(self, ofx_path: Path) -> dict:
        """
        Process OFX (Open Financial Exchange) files
        Standard bank statement format used by Brazilian banks:
        Itaú, Bradesco, Banco do Brasil, Santander, Caixa, Nubank
        """
        logger.info(f"📄 Processing OFX file: {ofx_path.name}")

        try:
            from ofxparse import OfxParser

            with open(ofx_path, "rb") as f:
                ofx = OfxParser.parse(f)

            transactions = []
            dates = []

            for account in ofx.accounts if hasattr(ofx, 'accounts') else [ofx.account]:
                for txn in account.statement.transactions:
                    amount = Decimal(str(txn.amount))
                    transaction_type = "receita" if amount > 0 else "despesa"

                    date_str = txn.date.strftime("%Y-%m-%d") if txn.date else None
                    if txn.date:
                        dates.append(txn.date)

                    transactions.append(Transaction(
                        date=date_str,
                        description=str(txn.memo or txn.payee or "")[:200],
                        category="nao_categorizado",
                        amount=abs(amount),
                        transaction_type=transaction_type,
                    ))

            # Build ledger
            total_income = sum(t.amount for t in transactions if t.transaction_type in ("income", "receita"))
            total_expense = sum(t.amount for t in transactions if t.transaction_type in ("expense", "despesa"))

            date_range = DateRangeSummary()
            if dates:
                date_range.start_date = min(dates).strftime("%Y-%m-%d")
                date_range.end_date = max(dates).strftime("%Y-%m-%d")
                date_range.total_days = (max(dates) - min(dates)).days

            ledger = TransactionLedger(
                file_name=ofx_path.name,
                total_transactions=len(transactions),
                date_range=date_range,
                total_income=total_income,
                total_expense=total_expense,
                net_balance=total_income - total_expense,
                by_category=[],
                transactions=transactions,
            )

            logger.info(f"✅ OFX processed: {len(transactions)} transactions")

            return {
                "file_name": ofx_path.name,
                "file_type": "ofx",
                "status": "success",
                "extracted_data": ledger,
            }

        except ImportError:
            error_msg = "ofxparse não está instalado. Execute: pip install ofxparse"
            logger.error(f"❌ {error_msg}")
            return {
                "file_name": ofx_path.name,
                "file_type": "ofx",
                "status": "error",
                "error": error_msg,
            }
        except Exception as e:
            logger.error(f"❌ OFX processing failed: {type(e).__name__}: {str(e)}")
            return {
                "file_name": ofx_path.name,
                "file_type": "ofx",
                "status": "error",
                "error": str(e),
            }

    def _process_ofc(self, ofc_path: Path) -> dict:
        """
        Process OFC (Open Financial Connectivity) files
        Legacy SGML-based bank statement format, predecessor to OFX
        Some older Brazilian banking systems still export this format
        """
        logger.info(f"📄 Processing OFC file: {ofc_path.name}")

        try:
            import re

            with open(ofc_path, "r", encoding="latin-1") as f:
                content = f.read()

            transactions = []
            dates = []

            # OFC uses SGML-like tags: <STMTTRN>, <DTPOSTED>, <TRNAMT>, <MEMO>
            # Parse transaction blocks
            txn_blocks = re.findall(
                r"<STMTTRN>(.*?)</STMTTRN>",
                content,
                re.DOTALL | re.IGNORECASE,
            )

            # If no STMTTRN blocks, try alternative format
            if not txn_blocks:
                txn_blocks = re.findall(
                    r"<STMTTRN>(.*?)(?=<STMTTRN>|</STMTTRNRS>|$)",
                    content,
                    re.DOTALL | re.IGNORECASE,
                )

            for block in txn_blocks:
                # Extract fields using regex
                date_match = re.search(r"<DTPOSTED>(\d{8})", block)
                amount_match = re.search(r"<TRNAMT>([-\d.,]+)", block)
                memo_match = re.search(r"<MEMO>(.+?)(?:<|\n)", block)
                name_match = re.search(r"<NAME>(.+?)(?:<|\n)", block)

                if amount_match:
                    try:
                        amount_str = amount_match.group(1).replace(",", ".")
                        amount = Decimal(amount_str)
                        transaction_type = "receita" if amount > 0 else "despesa"

                        date_str = None
                        if date_match:
                            raw_date = date_match.group(1)
                            date_str = f"{raw_date[0:4]}-{raw_date[4:6]}-{raw_date[6:8]}"
                            dates.append(datetime.strptime(date_str, "%Y-%m-%d"))

                        description = ""
                        if memo_match:
                            description = memo_match.group(1).strip()
                        elif name_match:
                            description = name_match.group(1).strip()

                        transactions.append(Transaction(
                            date=date_str,
                            description=description[:200],
                            category="nao_categorizado",
                            amount=abs(amount),
                            transaction_type=transaction_type,
                        ))
                    except (ValueError, ArithmeticError):
                        continue

            total_income = sum(t.amount for t in transactions if t.transaction_type in ("income", "receita"))
            total_expense = sum(t.amount for t in transactions if t.transaction_type in ("expense", "despesa"))

            date_range = DateRangeSummary()
            if dates:
                date_range.start_date = min(dates).strftime("%Y-%m-%d")
                date_range.end_date = max(dates).strftime("%Y-%m-%d")
                date_range.total_days = (max(dates) - min(dates)).days

            ledger = TransactionLedger(
                file_name=ofc_path.name,
                total_transactions=len(transactions),
                date_range=date_range,
                total_income=total_income,
                total_expense=total_expense,
                net_balance=total_income - total_expense,
                by_category=[],
                transactions=transactions,
            )

            logger.info(f"✅ OFC processed: {len(transactions)} transactions")

            return {
                "file_name": ofc_path.name,
                "file_type": "ofc",
                "status": "success",
                "extracted_data": ledger,
            }

        except Exception as e:
            logger.error(f"❌ OFC processing failed: {type(e).__name__}: {str(e)}")
            return {
                "file_name": ofc_path.name,
                "file_type": "ofc",
                "status": "error",
                "error": str(e),
            }

    def _process_docx(self, docx_path: Path) -> dict:
        """
        Process Word documents (.doc, .docx)
        Extracts text content and sends to AI for financial data extraction
        """
        logger.info(f"📄 Processing Word document: {docx_path.name}")

        try:
            from docx import Document

            doc = Document(str(docx_path))

            # Extract all text from paragraphs and tables
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n".join(text_parts)

            if not full_text.strip():
                return {
                    "file_name": docx_path.name,
                    "file_type": "docx",
                    "status": "error",
                    "error": "Documento Word vazio ou sem texto extraível",
                }

            logger.info(f"🤖 Calling AI for Word document extraction ({len(full_text)} chars)...")
            structured_data = self._extract_structured_data_from_excel(
                f"Arquivo Word: {docx_path.name}\n\n{full_text}"
            )

            logger.info(f"✅ Word document processed: {docx_path.name}")

            return {
                "file_name": docx_path.name,
                "file_type": "docx",
                "status": "success",
                "extracted_data": structured_data,
            }

        except ImportError:
            error_msg = "python-docx não está instalado. Execute: pip install python-docx"
            logger.error(f"❌ {error_msg}")
            return {
                "file_name": docx_path.name,
                "file_type": "docx",
                "status": "error",
                "error": error_msg,
            }
        except Exception as e:
            logger.error(f"❌ Word processing failed: {type(e).__name__}: {str(e)}")
            return {
                "file_name": docx_path.name,
                "file_type": "docx",
                "status": "error",
                "error": str(e),
            }

    def _process_txt(self, txt_path: Path) -> dict:
        """
        Process plain text files (.txt)
        Sends text content to AI for financial data extraction
        """
        logger.info(f"📄 Processing text file: {txt_path.name}")

        try:
            # Try multiple encodings common in Brazil
            text = None
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    with open(txt_path, "r", encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                return {
                    "file_name": txt_path.name,
                    "file_type": "txt",
                    "status": "error",
                    "error": "Não foi possível ler o arquivo de texto (encoding desconhecido)",
                }

            if not text.strip():
                return {
                    "file_name": txt_path.name,
                    "file_type": "txt",
                    "status": "error",
                    "error": "Arquivo de texto vazio",
                }

            logger.info(f"🤖 Calling AI for text file extraction ({len(text)} chars)...")
            structured_data = self._extract_structured_data_from_excel(
                f"Arquivo Texto: {txt_path.name}\n\n{text}"
            )

            logger.info(f"✅ Text file processed: {txt_path.name}")

            return {
                "file_name": txt_path.name,
                "file_type": "txt",
                "status": "success",
                "extracted_data": structured_data,
            }

        except Exception as e:
            logger.error(f"❌ Text processing failed: {type(e).__name__}: {str(e)}")
            return {
                "file_name": txt_path.name,
                "file_type": "txt",
                "status": "error",
                "error": str(e),
            }


def main():
    """Test the structured processor"""

    print("=== Structured Document Processor (Week 2) ===\n")

    processor = StructuredDocumentProcessor()
    print(f"Using AI Provider: {processor.ai_provider}\n")

    test_file = input("Enter path to test document: ").strip()

    if test_file:
        try:
            result = processor.process_document(test_file)

            print("\n" + "=" * 60)
            print("EXTRACTION RESULT:")
            print("=" * 60)
            print(f"File: {result['file_name']}")
            print(f"Type: {result['file_type']}")
            print(f"Status: {result['status']}")

            if result["status"] == "success":
                data = result["extracted_data"]
                print("\nStructured Data:")
                print("-" * 60)
                print(data.model_dump_json(indent=2))
            else:
                print(f"\nError: {result.get('error')}")

        except Exception as e:
            print(f"\nError: {e}")
            import traceback

            traceback.print_exc()


if __name__ == "__main__":
    main()
